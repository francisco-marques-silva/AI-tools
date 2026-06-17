"""
report_generator.py — Build Word (.docx) reports for the screening pipeline.

Produces N+1 documents per run:

  * `relatorio_geral_<ts>.docx` — the cross-project summary report
    (cover, methodology, data inventory, cost & metadata, the three Sens/Spec/F1
    trios, cost-effectiveness, workload, absolute efficiency, hours saved totals)

  * `relatorio_projeto_<safe-name>_<ts>.docx` — one self-contained report per
    project with the diagnostic confusion matrices, fulltext verification,
    test-retest binary matrices, error analyses, hours saved per execution
    and the per-project appendices (TIAB false positives, fulltext missed
    articles)

Chart PNGs are produced separately by `graphic.py` from the XLSX exported by
`chart_data.py`.
"""

import datetime
import re
from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .utils import (
    fmt, fmt_pct, normalise_model_name, load_file,
    compute_f1_lf, compute_metrics_vs_lf,
)
from .docx_helpers import (
    shade, set_cell, add_borders,
    add_heading, add_note, header_row,
)


# =====================================================================
#  Public orchestrator
# =====================================================================

def generate_report(projects, metadados, all_results, output_dir: Path):
    """Generate the cross-project report + one per-project report.

    Returns a list of pathlib.Path objects: [general, project_1, project_2, …].
    """
    ts_file = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = []

    # General (cross-project) report
    general_path = output_dir / f"relatorio_geral_{ts_file}.docx"
    _build_general_doc(projects, metadados, all_results).save(str(general_path))
    paths.append(general_path)

    # One report per project
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        safe = _safe_filename(proj["name"])
        proj_path = output_dir / f"relatorio_projeto_{safe}_{ts_file}.docx"
        _build_project_doc(
            pn, projects, metadados, all_results,
            general_filename=general_path.name,
        ).save(str(proj_path))
        paths.append(proj_path)

    return paths


# =====================================================================
#  Internal helpers
# =====================================================================

def _safe_filename(name: str) -> str:
    """Sanitize a project name for use in a filename."""
    s = re.sub(r"[^A-Za-z0-9_-]+", "_", str(name).strip().lower()).strip("_")
    return s or "project"


def _setup_doc():
    """Create a new Document with the report's default font."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    return doc


def _model_metrics_vs_lf(d, lf_res):
    """Sens/Spec/F1 of an AI test against the Listfinal gold standard."""
    if d is None:
        return None
    n_uni = d["n_paired"]
    n_pos = d["tp"] + d["fp"]
    return compute_metrics_vs_lf(n_uni, n_pos, lf_res)


def _human_metrics_vs_lf(hu_lf, pn):
    """Sens/Spec/F1 of the human reviewer against the Listfinal gold standard."""
    h = (hu_lf or {}).get(pn)
    if not h:
        return None
    return compute_metrics_vs_lf(h.get("n_universe"), h.get("n_human_pos"), h)


# =====================================================================
#  General (cross-project) document
# =====================================================================

def _build_general_doc(projects, metadados, all_results):
    """Cross-project summary report: cover, methodology and sections 1, 2, 5,
    7, 8, 9, 10, 11, 12, 13."""
    doc = _setup_doc()
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    table_counter = [0]

    def next_table():
        table_counter[0] += 1
        return table_counter[0]

    # Display name → project_norm key
    name_to_pn = {projects[pn]["name"]: pn for pn in projects}

    diag = all_results.get("diagnostic", {})
    ft = all_results.get("fulltext", {})
    lf = all_results.get("listfinal", {})
    tr = all_results.get("test_retest", {})
    hu_lf = all_results.get("human_listfinal", {})

    # ==================================================================
    #  COVER
    # ==================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("AI Screening — Cross-Project Analysis Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated on: {ts}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    proj_names = [projects[pn]["name"] for pn in sorted(projects.keys())]
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Projects analyzed: {', '.join(proj_names)}")
    run.font.size = Pt(10)

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note_p.add_run(
        "This file contains the cross-project summary tables. "
        "Project-specific detail (confusion matrices, fulltext verification, "
        "test-retest matrices and error appendices) is delivered in a "
        "separate per-project report for each project."
    )
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ==================================================================
    #  METHODOLOGICAL NOTES & REPORT GUIDE
    # ==================================================================
    add_heading(doc, "Methodological Notes & Report Guide", level=1)

    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(8)
    run = intro_p.add_run(
        "This report consolidates the cross-project analyses comparing AI "
        "screening decisions against human reviewers. The sections below are "
        "organized progressively: from data validation through diagnostic "
        "performance, error analysis, and cost-efficiency."
    )
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    section_guide = [
        ("Section 1 — Data Validation",
         "Inventory of all detected input files (AI results, human TIAB, Fulltext, "
         "Listfinal, metadata) and per-project completeness check."),
        ("Section 2 — Metadata and Costs",
         "Execution metadata table (model, parameters, time, tokens, cost) and "
         "cost summaries grouped by project and by model."),
        ("Section 5 — Listfinal Verification (Gold Standard)",
         "Capture rate of the final included articles per (project, model, test)."),
        ("Section 7 — False Negatives",
         "Articles included by the human reviewer but excluded by the AI."),
        ("Section 8 — False Positives",
         "Articles excluded by the human reviewer but included by the AI."),
        ("Section 9 — General Comparative Tables",
         "Three Sens/Spec/F1 trios — (A) models vs human TIAB, (B) models vs "
         "Listfinal gold standard, (C) human vs Listfinal — with deduplication "
         "audit. Test-retest Kappa and Cost are kept in a separate auxiliary table."),
        ("Section 10 — Cost-Effectiveness",
         "Cost (USD) vs mean sensitivity per model."),
        ("Section 11 — Workload Reduction",
         "Human screening time vs AI screening time per execution and per project."),
        ("Section 12 — Absolute Efficiency",
         "Combined selectivity + capture score. "
         "Efficiency Score = Listfinal Capture Rate × (1 − AI Positive Rate)."),
        ("Section 13 — Full-Text Hours Saved",
         "Hours of full-text reading the AI saves by including fewer articles at "
         "the TIAB stage. Assumes 2 reviewers × 5 minutes per article."),
    ]
    for title_text, desc in section_guide:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run = p.add_run(desc)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    doc.add_paragraph()
    meth_title = doc.add_paragraph()
    run = meth_title.add_run("Key Methodological Definitions")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    method_notes = [
        ("Binarization", "AI decisions (include, maybe, exclude) are binarized: "
         "include and maybe → positive (passes screening), exclude → negative."),
        ("Gold Standard Hierarchy", "Three reference levels: (1) Human TIAB decisions — "
         "intermediate reference; (2) Fulltext selection; (3) Listfinal — definitive "
         "gold standard used for sensitivity/specificity vs LF."),
        ("Cohen's Kappa", "Interpretation follows Landis & Koch (1977): < 0 Poor; "
         "0.00–0.20 Slight; 0.21–0.40 Fair; 0.41–0.60 Moderate; "
         "0.61–0.80 Substantial; 0.81–1.00 Almost Perfect."),
        ("Workload Reduction", "Human screening time (time_human) compared with AI "
         "screening time (time_ia). Speed factor = human time / AI time."),
        ("Absolute Efficiency", "Efficiency Score = LF Capture Rate × (1 − AI Positive Rate)."),
        ("Cost-Effectiveness", "Cost per sensitivity point = average cost / (sensitivity × 100)."),
        ("Deduplication", "Articles in Listfinal or Fulltext that do not appear in the "
         "TIAB universe are excluded from confusion matrices (see Section 9.5)."),
    ]
    for title_text, text in method_notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    doc.add_page_break()

    # ==================================================================
    #  SECTION 1 — DATA VALIDATION
    # ==================================================================
    add_heading(doc, "1. Data Validation", level=1)

    tn = next_table()
    add_heading(doc, f"Table {tn}. Detected Files Inventory", level=2)

    inventory_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        n_models = len(proj["models"])
        n_ai_files = sum(len(m["tests"]) for m in proj["models"].values())
        has_tiab = "Yes" if proj["human_tiab"] else "No"
        has_ft = "Yes" if proj["human_fulltext"] else "No"
        has_lf = "Yes" if proj.get("human_listfinal") else "No"
        lf_n = ""
        if proj.get("human_listfinal"):
            try:
                lf_df = load_file(str(proj["human_listfinal"]))
                lf_n = str(len(lf_df))
            except Exception:
                lf_n = "?"
        inventory_rows.append([
            proj["name"], str(n_models), str(n_ai_files),
            has_tiab, has_ft, has_lf, lf_n,
        ])

    inv_headers = ["Project", "Models", "AI Files", "Human TIAB",
                    "Human Fulltext", "Listfinal", "Listfinal N"]
    tbl = doc.add_table(rows=1 + len(inventory_rows), cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, inv_headers)
    for i, row_data in enumerate(inventory_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
            if j == 3 and val == "No":
                shade(tbl.cell(i + 1, j), "FFD6D6")
            if j in (4, 5) and val == "No":
                shade(tbl.cell(i + 1, j), "FFF3CD")

    doc.add_paragraph()

    validation_issues = all_results.get("validation_issues", [])
    if validation_issues:
        add_heading(doc, "Validation Notes", level=3)
        for issue in validation_issues:
            p = doc.add_paragraph(f"• {issue}")
            p.runs[0].font.size = Pt(9)
    else:
        p = doc.add_paragraph("✓ All files were validated successfully.")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()

    # ==================================================================
    #  SECTION 2 — METADATA AND COSTS
    # ==================================================================
    if metadados is not None and not metadados.empty:
        add_heading(doc, "2. Metadata and Costs", level=1)

        tn = next_table()
        add_heading(doc, f"Table {tn}. Execution Metadata", level=2)
        add_note(doc, "Details of each execution: model, parameters, time, tokens, and costs (USD).")

        meta_cols = ["Project", "Code", "Model", "Parameters", "Version",
                     "Time", "Tokens In", "Tokens Out", "Cost ($)"]
        meta_rows = []
        for _, row in metadados.iterrows():
            meta_rows.append([
                str(row.get("project", "")),
                str(row.get("code", "")),
                str(row.get("model", "")),
                str(row.get("parameter", "")),
                str(row.get("version", "")),
                str(row.get("time_ia", "")),
                str(int(row["tokens_input"])) if pd.notna(row.get("tokens_input")) else "-",
                str(int(row["tokens_output"])) if pd.notna(row.get("tokens_output")) else "-",
                f"{row['cost_total']:.2f}" if pd.notna(row.get("cost_total")) else "-",
            ])

        tbl = doc.add_table(rows=1 + len(meta_rows), cols=len(meta_cols))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, meta_cols)
        for i, row_data in enumerate(meta_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)

        doc.add_paragraph()

        # Cost Summary by Project
        tn = next_table()
        add_heading(doc, f"Table {tn}. Cost Summary by Project", level=2)

        meta_valid = metadados.dropna(subset=["project"])
        cost_summary = meta_valid.groupby("project").agg(
            n_execucoes=("cost_total", "count"),
            custo_total=("cost_total", "sum"),
            custo_medio=("cost_total", "mean"),
            tokens_in_total=("tokens_input", "sum"),
            tokens_out_total=("tokens_output", "sum"),
        ).reset_index()

        cost_headers = ["Project", "Executions", "Total Cost ($)", "Avg Cost ($)",
                        "Tokens In (total)", "Tokens Out (total)"]
        tbl = doc.add_table(rows=1 + len(cost_summary), cols=6)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cost_headers)
        for i, (_, row) in enumerate(cost_summary.iterrows()):
            vals = [
                str(row["project"]),
                str(int(row["n_execucoes"])),
                f"{row['custo_total']:.2f}",
                f"{row['custo_medio']:.2f}",
                str(int(row["tokens_in_total"])),
                str(int(row["tokens_out_total"])),
            ]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

        # Average Cost per Model
        tn = next_table()
        add_heading(doc, f"Table {tn}. Average Cost per Model", level=2)

        cost_model = metadados.dropna(subset=["project"]).copy()
        cost_model["model_norm"] = cost_model["model"].apply(normalise_model_name)
        cost_by_model = cost_model.groupby("model").agg(
            n=("cost_total", "count"),
            custo_medio=("cost_total", "mean"),
            custo_total=("cost_total", "sum"),
        ).reset_index().sort_values("custo_medio")

        cmod_headers = ["Model", "Executions", "Avg Cost ($)", "Total Cost ($)"]
        tbl = doc.add_table(rows=1 + len(cost_by_model), cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cmod_headers)
        for i, (_, row) in enumerate(cost_by_model.iterrows()):
            vals = [str(row["model"]), str(int(row["n"])),
                    f"{row['custo_medio']:.2f}", f"{row['custo_total']:.2f}"]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 5 — LISTFINAL VERIFICATION (combined)
    # ==================================================================
    add_heading(doc, "5. Listfinal Verification (True Gold Standard Capture Rate)", level=1)
    add_note(doc, "How many of the final included articles would have been retained "
             "by each AI model during TIAB screening.")

    if not lf:
        p = doc.add_paragraph("No project with Listfinal data available.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Listfinal Capture Rate — All Projects", level=2)

        lf_headers = ["Project", "Model", "Test", "Listfinal N", "Found",
                       "Captured", "Missed", "Capture Rate", "Miss Rate"]
        lf_rows = []
        for pn in sorted(lf.keys()):
            proj = projects[pn]
            for mn in sorted(lf[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(lf[pn][mn].keys()):
                    r = lf[pn][mn][test_num]
                    lf_rows.append([
                        proj["name"], model_name, f"{test_num}o",
                        str(r["n_listfinal"]), str(r["n_found"]),
                        str(r["n_captured"]), str(r["n_missed"]),
                        fmt_pct(r["capture_rate"]), fmt_pct(r["miss_rate"]),
                    ])

        tbl = doc.add_table(rows=1 + len(lf_rows), cols=len(lf_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, lf_headers)
        for i, row_data in enumerate(lf_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 7 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass
                if j == 6:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFF3CD")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 7 — FALSE NEGATIVES (count by model)
    # ==================================================================
    add_heading(doc, "7. False Negatives Analysis", level=1)
    add_note(doc, "Articles included by the human (maybe) but excluded by the AI (exclude). "
             "False negatives are the most critical errors in systematic review screening.")

    fn_results = all_results.get("false_negatives", {})
    if fn_results:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Negatives Count by Model", level=2)

        fn_headers = ["Project", "Model", "Test", "Total Paired", "False Neg.", "% of Total"]
        fn_rows = []
        for pn in sorted(fn_results.keys()):
            proj = projects[pn]
            for mn in sorted(fn_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(fn_results[pn][mn].keys()):
                    r = fn_results[pn][mn][test_num]
                    fn_count = r["fn"]
                    n_total = r["n_paired"]
                    fn_pct = fn_count / n_total * 100 if n_total > 0 else 0
                    fn_rows.append([
                        proj["name"], model_name, f"{test_num}º",
                        str(n_total), str(fn_count), f"{fn_pct:.1f}%",
                    ])

        tbl = doc.add_table(rows=1 + len(fn_rows), cols=len(fn_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fn_headers)
        for i, row_data in enumerate(fn_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 4:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        doc.add_paragraph()
    else:
        p = doc.add_paragraph("No false negative data (requires human reference).")
        p.runs[0].font.italic = True

    # ==================================================================
    #  SECTION 8 — FALSE POSITIVES (count by model)
    # ==================================================================
    add_heading(doc, "8. False Positives Analysis", level=1)
    add_note(doc, "Articles excluded by the human (exclude) but included by the AI (maybe).")

    fp_results = all_results.get("false_positives", {})
    if fp_results:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Positives Count by Model", level=2)

        fp_headers = ["Project", "Model", "Test", "Total Paired", "False Pos.", "% of Total",
                      "Human Excl. Articles", "FP Rate (of excl.)"]
        fp_rows = []
        for pn in sorted(fp_results.keys()):
            proj = projects[pn]
            for mn in sorted(fp_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(fp_results[pn][mn].keys()):
                    r = fp_results[pn][mn][test_num]
                    fp_count = r["fp"]
                    n_total = r["n_paired"]
                    n_excl_human = r["fp"] + r["tn"]
                    fp_pct = fp_count / n_total * 100 if n_total > 0 else 0
                    fp_rate = fp_count / n_excl_human * 100 if n_excl_human > 0 else 0
                    fp_rows.append([
                        proj["name"], model_name, f"{test_num}º",
                        str(n_total), str(fp_count), f"{fp_pct:.1f}%",
                        str(n_excl_human), f"{fp_rate:.1f}%",
                    ])

        tbl = doc.add_table(rows=1 + len(fp_rows), cols=len(fp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fp_headers)
        for i, row_data in enumerate(fp_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()
    else:
        p = doc.add_paragraph("No false positive data (requires human reference).")
        p.runs[0].font.italic = True

    doc.add_page_break()

    # ==================================================================
    #  SECTION 9 — GENERAL COMPARATIVE TABLES (3 Sens/Spec/F1 trios)
    # ==================================================================
    add_heading(doc, "9. General Comparative Tables", level=1)
    add_note(doc,
        "Performance summaries organised around three reference standards: "
        "(A) models vs human reviewer at TIAB; "
        "(B) models AND human reviewer vs the Listfinal gold standard; "
        "(C) combined view of the three Sens/Spec/F1 trios. "
        "All confusion matrices use the deduplicated paired universe; articles in "
        "Listfinal or Fulltext that do not appear in the TIAB universe are reported "
        "separately in subsection 9.5 and excluded from the metrics.")

    # ----- 9.1 — Models vs Human TIAB (Trio A) ------------------------
    add_heading(doc, "9.1 Models vs Human Reviewer (TIAB)", level=2)
    add_note(doc, "Trio A — how closely each model's TIAB decision matches the human "
             "reviewer's TIAB decision. Inclusion Rate = fraction of TIAB articles the "
             "model marked as include/maybe.")

    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. Trio A — Averaged across tests (principal view)", level=3)
    a_avg_headers = ["Project", "Model", "Sens. (TIAB)", "Spec. (TIAB)", "F1 (TIAB)", "Inclusion Rate"]
    a_avg_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            mi = proj["models"][mn]
            sens_l, spec_l, f1_l, inc_l = [], [], [], []
            for tn2 in sorted(mi["tests"].keys()):
                d = diag.get(pn, {}).get(mn, {}).get(tn2)
                if d:
                    m = d["metrics"]
                    if not np.isnan(m["Sensitivity"]): sens_l.append(m["Sensitivity"])
                    if not np.isnan(m["Specificity"]): spec_l.append(m["Specificity"])
                    if not np.isnan(m["F1 Score"]):    f1_l.append(m["F1 Score"])
                    n_paired = d["n_paired"]
                    if n_paired > 0:
                        inc_l.append((d["tp"] + d["fp"]) / n_paired)
            a_avg_rows.append([
                proj["name"], mi["name"],
                fmt_pct(np.mean(sens_l)) if sens_l else "-",
                fmt_pct(np.mean(spec_l)) if spec_l else "-",
                fmt(np.mean(f1_l), 3)    if f1_l    else "-",
                fmt_pct(np.mean(inc_l))  if inc_l   else "-",
            ])
    tbl = doc.add_table(rows=1 + len(a_avg_rows), cols=len(a_avg_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, a_avg_headers)
    for i, row_data in enumerate(a_avg_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

    hu_inc_lines = []
    for pn in sorted(projects.keys()):
        h = hu_lf.get(pn)
        if h and h.get("n_universe"):
            rate = h["n_human_pos"] / h["n_universe"]
            hu_inc_lines.append(
                f"{projects[pn]['name']}: {rate*100:.1f}% "
                f"({h['n_human_pos']}/{h['n_universe']})"
            )
    if hu_inc_lines:
        p = doc.add_paragraph()
        run = p.add_run("Human reviewer inclusion rate (TIAB) for context — "
                        + "  ·  ".join(hu_inc_lines))
        run.font.size = Pt(8)
        run.italic = True

    doc.add_paragraph()

    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. Trio A — Per test (secondary view)", level=3)
    a_pt_headers = ["Project", "Model", "Test", "Sens. (TIAB)", "Spec. (TIAB)",
                     "F1 (TIAB)", "Inclusion Rate"]
    a_pt_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            mi = proj["models"][mn]
            for tn2 in sorted(mi["tests"].keys()):
                d = diag.get(pn, {}).get(mn, {}).get(tn2)
                if not d:
                    a_pt_rows.append([proj["name"], mi["name"], f"{tn2}o",
                                       "-", "-", "-", "-"])
                    continue
                m = d["metrics"]
                n_paired = d["n_paired"]
                inc = (d["tp"] + d["fp"]) / n_paired if n_paired else float("nan")
                a_pt_rows.append([
                    proj["name"], mi["name"], f"{tn2}o",
                    fmt_pct(m["Sensitivity"]),
                    fmt_pct(m["Specificity"]),
                    fmt(m["F1 Score"], 3),
                    fmt_pct(inc),
                ])
    tbl = doc.add_table(rows=1 + len(a_pt_rows), cols=len(a_pt_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, a_pt_headers)
    for i, row_data in enumerate(a_pt_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
    doc.add_paragraph()

    # ----- 9.2 — Models AND Human vs Listfinal (Trios B + C) ----------
    add_heading(doc, "9.2 Models and Human Reviewer vs Listfinal (Gold Standard)", level=2)
    add_note(doc,
        "Trio B — models vs the Listfinal gold standard. Trio C — human reviewer vs "
        "the same gold standard, included as a baseline row per project. "
        "Sensitivity here = capture of the final included articles; Specificity = "
        "correct exclusion of non-final articles.")

    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. Trios B & C — Averaged across tests (principal view)", level=3)
    b_avg_headers = ["Project", "Entity", "Sens. (vs LF)", "Spec. (vs LF)", "F1 (vs LF)"]
    b_avg_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            mi = proj["models"][mn]
            sens_l, spec_l, f1_l = [], [], []
            for tn2 in sorted(mi["tests"].keys()):
                d = diag.get(pn, {}).get(mn, {}).get(tn2)
                lf_res = lf.get(pn, {}).get(mn, {}).get(tn2)
                mlf = _model_metrics_vs_lf(d, lf_res)
                if mlf:
                    if not np.isnan(mlf["sens_lf"]): sens_l.append(mlf["sens_lf"])
                    if not np.isnan(mlf["spec_lf"]): spec_l.append(mlf["spec_lf"])
                    if not np.isnan(mlf["f1_lf"]):   f1_l.append(mlf["f1_lf"])
            b_avg_rows.append([
                proj["name"], mi["name"],
                fmt_pct(np.mean(sens_l)) if sens_l else "-",
                fmt_pct(np.mean(spec_l)) if spec_l else "-",
                fmt(np.mean(f1_l), 3)    if f1_l    else "-",
            ])
        hlf = _human_metrics_vs_lf(hu_lf, pn)
        if hlf:
            b_avg_rows.append([
                proj["name"], "Human TIAB (baseline)",
                fmt_pct(hlf["sens_lf"]),
                fmt_pct(hlf["spec_lf"]),
                fmt(hlf["f1_lf"], 3),
            ])

    tbl = doc.add_table(rows=1 + len(b_avg_rows), cols=len(b_avg_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, b_avg_headers)
    for i, row_data in enumerate(b_avg_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
            if str(row_data[1]).startswith("Human"):
                shade(tbl.cell(i + 1, j), "FEF3C7")
    doc.add_paragraph()

    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. Trios B & C — Per test (secondary view)", level=3)
    b_pt_headers = ["Project", "Entity", "Test", "Sens. (vs LF)", "Spec. (vs LF)", "F1 (vs LF)"]
    b_pt_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            mi = proj["models"][mn]
            for tn2 in sorted(mi["tests"].keys()):
                d = diag.get(pn, {}).get(mn, {}).get(tn2)
                lf_res = lf.get(pn, {}).get(mn, {}).get(tn2)
                mlf = _model_metrics_vs_lf(d, lf_res)
                if mlf:
                    b_pt_rows.append([
                        proj["name"], mi["name"], f"{tn2}o",
                        fmt_pct(mlf["sens_lf"]),
                        fmt_pct(mlf["spec_lf"]),
                        fmt(mlf["f1_lf"], 3),
                    ])
                else:
                    b_pt_rows.append([proj["name"], mi["name"], f"{tn2}o",
                                       "-", "-", "-"])
        hlf = _human_metrics_vs_lf(hu_lf, pn)
        if hlf:
            b_pt_rows.append([
                proj["name"], "Human TIAB (baseline)", "—",
                fmt_pct(hlf["sens_lf"]),
                fmt_pct(hlf["spec_lf"]),
                fmt(hlf["f1_lf"], 3),
            ])

    tbl = doc.add_table(rows=1 + len(b_pt_rows), cols=len(b_pt_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, b_pt_headers)
    for i, row_data in enumerate(b_pt_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
            if str(row_data[1]).startswith("Human"):
                shade(tbl.cell(i + 1, j), "FEF3C7")
    doc.add_paragraph()

    # ----- 9.3 — Three trios combined --------------------------------
    add_heading(doc, "9.3 Three Trios Combined", level=2)
    add_note(doc, "All three Sens/Spec/F1 trios side-by-side. Model rows are "
             "averaged across tests. Human TIAB rows show Trio C (vs LF only) per project.")

    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. Comprehensive comparison — Sens/Spec/F1 across all references", level=3)
    c_headers = ["Project", "Entity",
                  "Sens (TIAB)", "Spec (TIAB)", "F1 (TIAB)",
                  "Sens (LF)", "Spec (LF)", "F1 (LF)"]
    c_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            mi = proj["models"][mn]
            s_t, sp_t, f_t = [], [], []
            s_l, sp_l, f_l = [], [], []
            for tn2 in sorted(mi["tests"].keys()):
                d = diag.get(pn, {}).get(mn, {}).get(tn2)
                lf_res = lf.get(pn, {}).get(mn, {}).get(tn2)
                if d:
                    m = d["metrics"]
                    if not np.isnan(m["Sensitivity"]): s_t.append(m["Sensitivity"])
                    if not np.isnan(m["Specificity"]): sp_t.append(m["Specificity"])
                    if not np.isnan(m["F1 Score"]):    f_t.append(m["F1 Score"])
                mlf = _model_metrics_vs_lf(d, lf_res)
                if mlf:
                    if not np.isnan(mlf["sens_lf"]): s_l.append(mlf["sens_lf"])
                    if not np.isnan(mlf["spec_lf"]): sp_l.append(mlf["spec_lf"])
                    if not np.isnan(mlf["f1_lf"]):   f_l.append(mlf["f1_lf"])
            c_rows.append([
                proj["name"], mi["name"],
                fmt_pct(np.mean(s_t)) if s_t else "-",
                fmt_pct(np.mean(sp_t)) if sp_t else "-",
                fmt(np.mean(f_t), 3)   if f_t   else "-",
                fmt_pct(np.mean(s_l)) if s_l else "-",
                fmt_pct(np.mean(sp_l)) if sp_l else "-",
                fmt(np.mean(f_l), 3)   if f_l   else "-",
            ])
        hlf = _human_metrics_vs_lf(hu_lf, pn)
        if hlf:
            c_rows.append([
                proj["name"], "Human TIAB (baseline)",
                "—", "—", "—",
                fmt_pct(hlf["sens_lf"]),
                fmt_pct(hlf["spec_lf"]),
                fmt(hlf["f1_lf"], 3),
            ])

    tbl = doc.add_table(rows=1 + len(c_rows), cols=len(c_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, c_headers)
    for i, row_data in enumerate(c_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
            if str(row_data[1]).startswith("Human"):
                shade(tbl.cell(i + 1, j), "FEF3C7")
    doc.add_paragraph()

    # ----- 9.4 — Reproducibility & Cost (auxiliary) -------------------
    add_heading(doc, "9.4 Reproducibility & Cost (Auxiliary)", level=2)
    add_note(doc, "Test-retest Kappa and total cost per (project, model). Kept here "
             "for reference; they are not part of the Sens/Spec/F1 comparison.")

    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. Test-Retest Kappa and Cost", level=3)
    aux_headers = ["Project", "Model", "Kappa (T-R)", "Cost ($, total)"]
    aux_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            mi = proj["models"][mn]
            tr_res = tr.get(pn, {}).get(mn)
            kappa_tr = fmt(tr_res["kappa"], 3) if tr_res else "-"
            cost_vals = []
            if metadados is not None:
                for tn2 in sorted(mi["tests"].keys()):
                    code = mi["tests"][tn2]["code"]
                    meta_m = metadados[metadados["code"].astype(str) == str(code)]
                    if not meta_m.empty and pd.notna(meta_m.iloc[0]["cost_total"]):
                        cost_vals.append(meta_m.iloc[0]["cost_total"])
            cost = fmt(sum(cost_vals), 2) if cost_vals else "-"
            aux_rows.append([proj["name"], mi["name"], kappa_tr, cost])

    tbl = doc.add_table(rows=1 + len(aux_rows), cols=len(aux_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, aux_headers)
    for i, row_data in enumerate(aux_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
    doc.add_paragraph()

    # ----- 9.5 — Deduplication audit ----------------------------------
    add_heading(doc, "9.5 Deduplication Audit", level=2)
    add_note(doc, "Articles from Listfinal or Fulltext that did not appear in the TIAB "
             "universe (possible deduplication artifacts). These are excluded from the "
             "confusion matrices above so that no entity is penalised for articles it "
             "never had a chance to screen.")

    audit_headers = ["Project", "TIAB N (paired)",
                      "LF total", "LF not in TIAB",
                      "FT total", "FT not in TIAB"]
    audit_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        n_paired = "-"
        pd_diag = diag.get(pn) or {}
        for mn2 in sorted(pd_diag.keys()):
            for tn2 in sorted(pd_diag[mn2].keys()):
                dd = pd_diag[mn2][tn2]
                if dd is not None:
                    n_paired = str(dd["n_paired"])
                    break
            if n_paired != "-":
                break

        lf_total, lf_missing = "-", "-"
        hu = hu_lf.get(pn)
        if hu:
            lf_total = str(hu["n_listfinal"])
            lf_missing = str(hu["n_not_found"])
        else:
            pd_lf = lf.get(pn) or {}
            for mn2 in sorted(pd_lf.keys()):
                for tn2 in sorted(pd_lf[mn2].keys()):
                    rr = pd_lf[mn2][tn2]
                    lf_total = str(rr["n_listfinal"])
                    lf_missing = str(rr["n_not_found"])
                    break
                if lf_total != "-":
                    break

        ft_total, ft_missing = "-", "-"
        pd_ft = ft.get(pn) or {}
        for mn2 in sorted(pd_ft.keys()):
            for tn2 in sorted(pd_ft[mn2].keys()):
                rr = pd_ft[mn2][tn2]
                ft_total = str(rr["n_fulltext"])
                ft_missing = str(rr["n_not_found"])
                break
            if ft_total != "-":
                break

        audit_rows.append([proj["name"], n_paired,
                            lf_total, lf_missing,
                            ft_total, ft_missing])

    tbl = doc.add_table(rows=1 + len(audit_rows), cols=len(audit_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, audit_headers)
    for i, row_data in enumerate(audit_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
            if j in (3, 5) and val not in ("-", "0"):
                try:
                    if int(val) > 0:
                        shade(tbl.cell(i + 1, j), "FFD6D6")
                except ValueError:
                    pass
    doc.add_paragraph()

    # ==================================================================
    #  SECTION 10 — COST-EFFECTIVENESS
    # ==================================================================
    if metadados is not None and diag:
        add_heading(doc, "10. Cost-Effectiveness Analysis", level=1)
        add_note(doc, "Relationship between cost (USD) and diagnostic performance.")

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Cost vs. Sensitivity per Model (test average)", level=2)

        cost_eff_headers = ["Model", "Project", "Avg Sens.", "Avg Spec.",
                            "Avg F1 (LF)", "Avg Cost ($)", "Cost per Sens. point"]
        cost_eff_rows = []

        for pn in sorted(diag.keys()):
            proj = projects[pn]
            for mn in sorted(diag[pn].keys()):
                model_name = proj["models"][mn]["name"]
                tests = diag[pn][mn]
                sens_vals, spec_vals, f1_vals, cost_vals = [], [], [], []
                for tn2, r in tests.items():
                    if r is None:
                        continue
                    s = r["metrics"]["Sensitivity"]
                    if not np.isnan(s):
                        sens_vals.append(s)
                    sp = r["metrics"]["Specificity"]
                    if not np.isnan(sp):
                        spec_vals.append(sp)
                    lf_r = lf.get(pn, {}).get(mn, {}).get(tn2)
                    f1_lf = compute_f1_lf(r, lf_r)
                    if not np.isnan(f1_lf):
                        f1_vals.append(f1_lf)
                    code = proj["models"][mn]["tests"][tn2]["code"]
                    meta_m = metadados[metadados["code"].astype(str) == str(code)]
                    if not meta_m.empty and pd.notna(meta_m.iloc[0]["cost_total"]):
                        cost_vals.append(meta_m.iloc[0]["cost_total"])

                avg_sens = np.mean(sens_vals) if sens_vals else float("nan")
                avg_spec = np.mean(spec_vals) if spec_vals else float("nan")
                avg_f1 = np.mean(f1_vals) if f1_vals else float("nan")
                avg_cost = np.mean(cost_vals) if cost_vals else float("nan")
                cost_per_sens = avg_cost / (avg_sens * 100) if (
                    not np.isnan(avg_cost) and not np.isnan(avg_sens) and avg_sens > 0
                ) else float("nan")

                cost_eff_rows.append([
                    model_name, proj["name"],
                    fmt_pct(avg_sens), fmt_pct(avg_spec),
                    fmt(avg_f1, 3),
                    fmt(avg_cost, 2) if not np.isnan(avg_cost) else "-",
                    fmt(cost_per_sens, 3) if not np.isnan(cost_per_sens) else "-",
                ])

        tbl = doc.add_table(rows=1 + len(cost_eff_rows), cols=len(cost_eff_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cost_eff_headers)
        for i, row_data in enumerate(cost_eff_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()

    # ==================================================================
    #  SECTION 11 — WORKLOAD REDUCTION
    # ==================================================================
    if metadados is not None:
        add_heading(doc, "11. Workload Reduction Analysis", level=1)
        add_note(doc, "Compares human screening time (time_human) with AI screening time (time_ia).")

        def parse_td(val):
            if pd.isna(val):
                return float("nan")
            if isinstance(val, pd.Timedelta):
                return val.total_seconds() / 3600.0
            s = str(val).strip()
            if not s:
                return float("nan")
            try:
                return pd.to_timedelta(s).total_seconds() / 3600.0
            except Exception:
                return float("nan")

        def fmt_hours(h):
            if np.isnan(h):
                return "-"
            hr = int(h)
            mn = int((h - hr) * 60)
            return f"{hr}h {mn:02d}m" if hr > 0 else f"{mn}m"

        meta_work = metadados.copy()
        meta_work["_h_human"] = meta_work["time_human"].apply(parse_td)
        meta_work["_h_ia"] = meta_work["time_ia"].apply(parse_td)

        def _proj_key(val):
            s = str(val).strip()
            return name_to_pn.get(s, s.lower())

        meta_work["_proj_key"] = meta_work["project"].apply(_proj_key)

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Workload Reduction — Per Execution", level=2)

        wr_headers = ["Project", "Model", "Code", "Human Time", "AI Time",
                      "Time Saved", "Reduction (%)", "Speed Factor"]
        wr_rows = []
        for _, row in meta_work.iterrows():
            h_human = row["_h_human"]
            h_ia = row["_h_ia"]
            saved = h_human - h_ia if not (np.isnan(h_human) or np.isnan(h_ia)) else float("nan")
            reduction = (saved / h_human * 100) if (not np.isnan(saved) and h_human > 0) else float("nan")
            factor = (h_human / h_ia) if (not np.isnan(h_human) and not np.isnan(h_ia) and h_ia > 0) else float("nan")
            wr_rows.append([
                str(row.get("project", "")),
                str(row.get("model", "")),
                str(row.get("code", "")),
                fmt_hours(h_human),
                fmt_hours(h_ia),
                fmt_hours(saved),
                f"{reduction:.1f}%" if not np.isnan(reduction) else "-",
                f"{factor:.0f}x" if not np.isnan(factor) else "-",
            ])

        tbl = doc.add_table(rows=1 + len(wr_rows), cols=len(wr_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, wr_headers)
        for i, row_data in enumerate(wr_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 6 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 90:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val >= 50:
                            shade(tbl.cell(i + 1, j), "E8F5E9")
                    except ValueError:
                        pass

        doc.add_paragraph()

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Workload Reduction Summary by Project", level=2)

        wr_proj = meta_work.groupby("_proj_key").agg(
            display_name=("project", "first"),
            human_hours=("_h_human", "first"),
            avg_ia_hours=("_h_ia", "mean"),
            min_ia_hours=("_h_ia", "min"),
            max_ia_hours=("_h_ia", "max"),
            n_runs=("_h_ia", "count"),
        ).reset_index()

        wp_headers = ["Project", "Total Articles", "Human Time", "Avg AI Time", "Fastest AI",
                      "Avg Reduction (%)", "Avg Speed Factor"]
        tbl = doc.add_table(rows=1 + len(wr_proj), cols=len(wp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, wp_headers)
        for i, (_, row) in enumerate(wr_proj.iterrows()):
            pn_norm = row["_proj_key"]
            display_name = row["display_name"]
            if pn_norm in projects:
                display_name = projects[pn_norm]["name"]
            h_h = row["human_hours"]
            avg_ia = row["avg_ia_hours"]
            fast = row["min_ia_hours"]
            avg_red = ((h_h - avg_ia) / h_h * 100) if (not np.isnan(h_h) and not np.isnan(avg_ia) and h_h > 0) else float("nan")
            avg_fac = (h_h / avg_ia) if (not np.isnan(h_h) and not np.isnan(avg_ia) and avg_ia > 0) else float("nan")
            total_arts = "-"
            proj_d = diag.get(pn_norm) or {}
            for mn2 in sorted(proj_d.keys()):
                for tn2 in sorted(proj_d[mn2].keys()):
                    d2 = proj_d[mn2][tn2]
                    if d2 is not None:
                        total_arts = str(d2["n_paired"])
                        break
                if total_arts != "-":
                    break
            vals = [
                str(display_name),
                total_arts,
                fmt_hours(h_h),
                fmt_hours(avg_ia),
                fmt_hours(fast),
                f"{avg_red:.1f}%" if not np.isnan(avg_red) else "-",
                f"{avg_fac:.0f}x" if not np.isnan(avg_fac) else "-",
            ]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

    # ==================================================================
    #  SECTION 12 — ABSOLUTE EFFICIENCY
    # ==================================================================
    if diag and lf:
        add_heading(doc, "12. Absolute Efficiency Analysis (TIAB Reduction vs Listfinal Retention)", level=1)
        add_note(doc, "Measures how much the AI reduces the TIAB workload while still "
                 "retaining truly relevant articles.")

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Absolute Efficiency — TIAB Selection Volume vs Listfinal Retention", level=2)

        eff_headers = ["Project", "Model", "Test", "TIAB N",
                       "AI Positives", "AI Pos. (%)", "Human Positives", "Human Pos. (%)",
                       "Reduction", "Capture Rate", "Efficiency Score"]
        eff_rows = []

        for pn in sorted(projects.keys()):
            proj = projects[pn]
            proj_diag = diag.get(pn, {})
            proj_lf = lf.get(pn, {})

            for mn in sorted(proj["models"].keys()):
                model_info = proj["models"][mn]
                model_name = model_info["name"]
                for test_num in sorted(model_info["tests"].keys()):
                    d = proj_diag.get(mn, {}).get(test_num)
                    l = proj_lf.get(mn, {}).get(test_num)
                    if not d:
                        continue

                    n_total = d["n_paired"]
                    ai_pos = d["tp"] + d["fp"]
                    hu_pos = d["tp"] + d["fn"]
                    ai_pct = ai_pos / n_total * 100 if n_total > 0 else float("nan")
                    hu_pct = hu_pos / n_total * 100 if n_total > 0 else float("nan")
                    reduction = ((hu_pos - ai_pos) / hu_pos * 100) if hu_pos > 0 else float("nan")

                    lf_cap = "-"
                    eff_score = "-"
                    if l:
                        cap_rate = l["capture_rate"]
                        lf_cap = fmt_pct(cap_rate)
                        if not np.isnan(cap_rate) and not np.isnan(ai_pct):
                            score = cap_rate * (1 - ai_pct / 100)
                            eff_score = fmt(score, 3)

                    eff_rows.append([
                        proj["name"], model_name, f"{test_num}o",
                        str(n_total),
                        str(ai_pos), f"{ai_pct:.1f}%" if not np.isnan(ai_pct) else "-",
                        str(hu_pos), f"{hu_pct:.1f}%" if not np.isnan(hu_pct) else "-",
                        f"{reduction:+.1f}%" if not np.isnan(reduction) else "-",
                        lf_cap, eff_score,
                    ])

        if eff_rows:
            tbl = doc.add_table(rows=1 + len(eff_rows), cols=len(eff_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(tbl)
            header_row(tbl, eff_headers)
            for i, row_data in enumerate(eff_rows):
                for j, val in enumerate(row_data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
                    if j == 9 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val >= 95:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val < 80:
                                shade(tbl.cell(i + 1, j), "FFD6D6")
                        except ValueError:
                            pass
                    if j == 5 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val < 30:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val > 70:
                                shade(tbl.cell(i + 1, j), "FFF3CD")
                        except ValueError:
                            pass

            doc.add_paragraph()
            add_note(doc, "AI Positives = articles the AI selected for further review. "
                     "Reduction = % fewer articles selected by AI vs human. Negative = AI selected more. "
                     "Efficiency Score = LF Capture Rate × (1 − AI Positive Rate). Higher = better.")
        else:
            p = doc.add_paragraph("Insufficient data for absolute efficiency analysis.")
            p.runs[0].font.italic = True

        doc.add_paragraph()

    # ==================================================================
    #  SECTION 13 — FULL-TEXT HOURS SAVED (totals + per project)
    # ==================================================================
    if diag:
        add_heading(doc, "13. Full-Text Reading Hours Saved", level=1)
        add_note(doc, "Hours of full-text reading saved because the AI included fewer "
                 "articles at TIAB than the human. Assumes 2 reviewers × 5 minutes per "
                 "article. Per-execution detail is shown in each project's report. "
                 "Negative values mean the AI was more inclusive than the human (extra workload).")

        # Aggregated total per model
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Total Full-Text Hours Saved per Model (across projects)", level=2)

        model_totals = {}
        for pn in sorted(projects.keys()):
            proj = projects[pn]
            for mn in sorted(proj["models"].keys()):
                model_name = proj["models"][mn]["name"]
                if model_name not in model_totals:
                    model_totals[model_name] = {"arts": 0, "hrs": 0.0, "runs": 0}
                for test_num in sorted(proj["models"][mn]["tests"].keys()):
                    d = diag.get(pn, {}).get(mn, {}).get(test_num)
                    if not d:
                        continue
                    ai_inc = d["tp"] + d["fp"]
                    hu_inc = d["tp"] + d["fn"]
                    arts_saved = hu_inc - ai_inc
                    model_totals[model_name]["arts"] += arts_saved
                    model_totals[model_name]["hrs"] += arts_saved * 2 * 5 / 60.0
                    model_totals[model_name]["runs"] += 1

        tot_headers = ["Model", "Runs", "Total Articles Saved", "Total Hours Saved"]
        tot_rows = sorted(model_totals.items(),
                          key=lambda kv: kv[1]["hrs"], reverse=True)
        tbl = doc.add_table(rows=1 + len(tot_rows), cols=len(tot_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, tot_headers)
        for i, (mname, vals) in enumerate(tot_rows):
            row_vals = [
                mname,
                str(vals["runs"]),
                f"{vals['arts']:+d}",
                f"{vals['hrs']:+.1f}h",
            ]
            for j, v in enumerate(row_vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)
                if j == 3:
                    if vals["hrs"] > 0:
                        shade(tbl.cell(i + 1, j), "D5F5E3")
                    elif vals["hrs"] < 0:
                        shade(tbl.cell(i + 1, j), "FFD6D6")

        doc.add_paragraph()

        # Aggregated per project (summary across models)
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Hours Saved Summary by Project", level=2)

        proj_totals_headers = ["Project", "Best Model", "Best Hours Saved",
                                "Avg Hours Saved across Models"]
        proj_rows = []
        for pn in sorted(projects.keys()):
            proj = projects[pn]
            model_hours = {}
            for mn in sorted(proj["models"].keys()):
                model_name = proj["models"][mn]["name"]
                hrs = 0.0
                cnt = 0
                for tn2 in sorted(proj["models"][mn]["tests"].keys()):
                    d = diag.get(pn, {}).get(mn, {}).get(tn2)
                    if not d:
                        continue
                    arts = (d["tp"] + d["fn"]) - (d["tp"] + d["fp"])
                    hrs += arts * 2 * 5 / 60.0
                    cnt += 1
                if cnt > 0:
                    model_hours[model_name] = hrs / cnt
            if not model_hours:
                continue
            best_model, best_hours = max(model_hours.items(), key=lambda kv: kv[1])
            avg_hours = sum(model_hours.values()) / len(model_hours)
            proj_rows.append([
                proj["name"], best_model,
                f"{best_hours:+.1f}h",
                f"{avg_hours:+.1f}h",
            ])

        if proj_rows:
            tbl = doc.add_table(rows=1 + len(proj_rows), cols=len(proj_totals_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(tbl)
            header_row(tbl, proj_totals_headers)
            for i, row_data in enumerate(proj_rows):
                for j, val in enumerate(row_data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()

    return doc


# =====================================================================
#  Per-project document
# =====================================================================

def _build_project_doc(pn, projects, metadados, all_results, general_filename=""):
    """Self-contained per-project report: diagnostic confusion matrices,
    fulltext verification, listfinal, test-retest matrices, FN/FP detail,
    hours saved per execution, and per-project appendices A and B."""
    doc = _setup_doc()
    proj = projects[pn]
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    table_counter = [0]

    def next_table():
        table_counter[0] += 1
        return table_counter[0]

    diag = all_results.get("diagnostic", {})
    ft = all_results.get("fulltext", {})
    lf = all_results.get("listfinal", {})
    tr = all_results.get("test_retest", {})
    hu_lf = all_results.get("human_listfinal", {})
    fn_results = all_results.get("false_negatives", {})
    fp_results = all_results.get("false_positives", {})

    proj_diag = diag.get(pn, {})
    proj_ft = ft.get(pn, {})
    proj_lf = lf.get(pn, {})
    proj_tr = tr.get(pn, {})
    proj_fn = fn_results.get(pn, {})
    proj_fp = fp_results.get(pn, {})

    # ==================================================================
    #  COVER
    # ==================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"AI Screening — Project Report")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj_p.add_run(proj["name"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 73, 125)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated on: {ts}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    if general_filename:
        sub2 = doc.add_paragraph()
        sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub2.add_run(f"Cross-project summary: {general_filename}")
        run.font.size = Pt(8.5)
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Brief data inventory
    inv_p = doc.add_paragraph()
    inv_p.paragraph_format.space_before = Pt(18)
    inv_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n_models = len(proj["models"])
    n_ai_files = sum(len(m["tests"]) for m in proj["models"].values())
    has_tiab = "Yes" if proj["human_tiab"] else "No"
    has_ft = "Yes" if proj["human_fulltext"] else "No"
    has_lf = "Yes" if proj.get("human_listfinal") else "No"
    inv_text = (
        f"Models: {n_models}  |  AI executions: {n_ai_files}  |  "
        f"Human TIAB: {has_tiab}  |  Human Fulltext: {has_ft}  |  Listfinal: {has_lf}"
    )
    run = inv_p.add_run(inv_text)
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    # ==================================================================
    #  1. TIAB AGREEMENT ANALYSIS (AI vs Human TIAB)
    # ==================================================================
    add_heading(doc, "1. TIAB Agreement Analysis (AI vs Human Screener)", level=1)
    add_note(doc, "Comparison between AI and human screener decisions at the TIAB level. "
             "Note: the human TIAB decision is an intermediate reference, not the final gold standard.")

    if not proj_diag:
        p = doc.add_paragraph("No human TIAB reference for this project.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Model Comparison", level=2)
        add_note(doc, "Diagnostic metrics for each model and test, compared to the human reviewer.")

        comp_headers = ["Model", "Test", "N", "TP", "FP", "FN", "TN",
                        "Sens.", "Spec.", "PPV", "NPV", "Acc.", "F1", "Kappa", "95% CI", "Interpretation"]
        comp_rows = []
        for mn in sorted(proj_diag.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_diag[mn].keys()):
                r = proj_diag[mn][test_num]
                if r is None:
                    continue
                m = r["metrics"]
                comp_rows.append([
                    model_name, f"{test_num}º",
                    str(r["n_paired"]),
                    str(r["tp"]), str(r["fp"]), str(r["fn"]), str(r["tn"]),
                    fmt_pct(m["Sensitivity"]),
                    fmt_pct(m["Specificity"]),
                    fmt_pct(m["PPV (Precision)"]),
                    fmt_pct(m["NPV"]),
                    fmt_pct(m["Accuracy"]),
                    fmt(m["F1 Score"], 3),
                    fmt(r["kappa"], 3),
                    f"[{fmt(r['kappa_ci_lo'], 2)}, {fmt(r['kappa_ci_hi'], 2)}]",
                    r["kappa_interp"],
                ])

        tbl = doc.add_table(rows=1 + len(comp_rows), cols=len(comp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, comp_headers)
        for i, row_data in enumerate(comp_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
                if j == 7 and val not in ("-", "N/A"):
                    try:
                        v = float(val.replace("%", ""))
                        if v >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif v < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass

        doc.add_paragraph()

        # Confusion matrices
        for mn in sorted(proj_diag.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_diag[mn].keys()):
                r = proj_diag[mn][test_num]
                if r is None:
                    continue
                tn_num = next_table()
                add_heading(doc, f"Table {tn_num}. Confusion Matrix — {model_name} (Test {test_num})", level=3)

                cm_tbl = doc.add_table(rows=4, cols=4)
                cm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                add_borders(cm_tbl)

                cm_headers = ["", "Human: Maybe", "Human: Exclude", "Total"]
                for j, h in enumerate(cm_headers):
                    set_cell(cm_tbl.cell(0, j), h, bold=True, font_size=Pt(8))
                    shade(cm_tbl.cell(0, j), "D9E2F3")

                cm_data = [
                    ("AI: Maybe",   str(r["tp"]), str(r["fp"]), str(r["tp"] + r["fp"])),
                    ("AI: Exclude", str(r["fn"]), str(r["tn"]), str(r["fn"] + r["tn"])),
                    ("Total",       str(r["tp"] + r["fn"]), str(r["fp"] + r["tn"]),
                     str(r["tp"] + r["fp"] + r["fn"] + r["tn"])),
                ]
                for i, (label, *vals) in enumerate(cm_data, start=1):
                    set_cell(cm_tbl.cell(i, 0), label, bold=True,
                             font_size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT)
                    for j2, v in enumerate(vals, start=1):
                        set_cell(cm_tbl.cell(i, j2), v, font_size=Pt(8))
                    if i == 1:
                        shade(cm_tbl.cell(i, 1), "D5F5E3")
                    if i == 2:
                        shade(cm_tbl.cell(i, 2), "D5F5E3")

                doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  2. FULLTEXT VERIFICATION
    # ==================================================================
    add_heading(doc, "2. Fulltext Verification (Capture Rate)", level=1)
    add_note(doc, "Checks whether articles included in the human full-text review would "
             "have been retained by the AI during TIAB screening.")

    if not proj_ft:
        p = doc.add_paragraph("No fulltext data available for this project.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Fulltext Capture Rate", level=2)

        ft_headers = ["Model", "Test", "FT Articles", "Found",
                      "Captured", "Missed", "Capture Rate", "Miss Rate"]
        ft_rows = []
        for mn in sorted(proj_ft.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_ft[mn].keys()):
                r = proj_ft[mn][test_num]
                ft_rows.append([
                    model_name, f"{test_num}º",
                    str(r["n_fulltext"]), str(r["n_found"]),
                    str(r["n_captured"]), str(r["n_missed"]),
                    fmt_pct(r["capture_rate"]), fmt_pct(r["miss_rate"]),
                ])

        tbl = doc.add_table(rows=1 + len(ft_rows), cols=len(ft_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, ft_headers)
        for i, row_data in enumerate(ft_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 6 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass
                if j == 5:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFF3CD")
                    except ValueError:
                        pass

        doc.add_paragraph()

        all_missed = set()
        for mn in sorted(proj_ft.keys()):
            for test_num in sorted(proj_ft[mn].keys()):
                r = proj_ft[mn][test_num]
                for t in r.get("missed_titles", []):
                    all_missed.add(t)
        if all_missed:
            p = doc.add_paragraph(
                f"Note: {len(all_missed)} unique article(s) were missed by at least one AI run. "
                "See Appendix B for details."
            )
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True

    doc.add_page_break()

    # ==================================================================
    #  3. LISTFINAL VERIFICATION
    # ==================================================================
    add_heading(doc, "3. Listfinal Verification (Gold Standard Capture Rate)", level=1)
    add_note(doc, "Capture of the final included articles (post full-text reading). "
             "This is the definitive performance measure.")

    if not proj_lf:
        p = doc.add_paragraph("No Listfinal data available for this project.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Listfinal Capture Rate", level=2)

        lf_headers = ["Model", "Test", "Listfinal N", "Found",
                       "Captured", "Missed", "Capture Rate", "Miss Rate"]
        lf_rows = []
        for mn in sorted(proj_lf.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_lf[mn].keys()):
                r = proj_lf[mn][test_num]
                lf_rows.append([
                    model_name, f"{test_num}o",
                    str(r["n_listfinal"]), str(r["n_found"]),
                    str(r["n_captured"]), str(r["n_missed"]),
                    fmt_pct(r["capture_rate"]), fmt_pct(r["miss_rate"]),
                ])

        tbl = doc.add_table(rows=1 + len(lf_rows), cols=len(lf_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, lf_headers)
        for i, row_data in enumerate(lf_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 6 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass
                if j == 5:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFF3CD")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        # Human baseline vs LF for this project, if available
        hlf = _human_metrics_vs_lf(hu_lf, pn)
        if hlf:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(
                f"Human reviewer baseline vs Listfinal — "
                f"Sens: {fmt_pct(hlf['sens_lf'])}, "
                f"Spec: {fmt_pct(hlf['spec_lf'])}, "
                f"F1: {fmt(hlf['f1_lf'], 3)} "
                f"(N universe: {hlf['n_universe']}, "
                f"human positives: {hlf['n_positives']})."
            )
            run.font.size = Pt(9)
            run.italic = True

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  4. TEST-RETEST (per project detail)
    # ==================================================================
    add_heading(doc, "4. Test-Retest (Reproducibility)", level=1)
    add_note(doc, "Compares two runs of the same model on the same dataset.")

    if not proj_tr:
        p = doc.add_paragraph("No test-retest pairs available for this project.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Test-Retest Summary", level=2)
        tr_headers = ["Model", "N", "Exact Agree.", "Binary Agree.",
                       "Discordant", "Kappa", "95% CI", "Interpretation"]
        tr_rows = []
        for mn in sorted(proj_tr.keys()):
            model_name = proj["models"][mn]["name"]
            r = proj_tr[mn]
            tr_rows.append([
                model_name,
                str(r["n_total"]),
                f"{r['exact_match']} ({fmt_pct(r['exact_pct'])})",
                f"{r['binary_match']} ({fmt_pct(r['binary_pct'])})",
                str(r["n_discordant"]),
                fmt(r["kappa"], 3),
                f"[{fmt(r['kappa_ci_lo'], 2)}, {fmt(r['kappa_ci_hi'], 2)}]",
                r["kappa_interp"],
            ])

        tbl = doc.add_table(rows=1 + len(tr_rows), cols=len(tr_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, tr_headers)
        for i, row_data in enumerate(tr_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()

        # Per-model binary matrices
        for mn in sorted(proj_tr.keys()):
            model_name = proj["models"][mn]["name"]
            r = proj_tr[mn]

            tn_num = next_table()
            add_heading(doc, f"Table {tn_num}. Binary Test-Retest Matrix — {model_name}", level=3)

            cm_tbl = doc.add_table(rows=4, cols=4)
            cm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(cm_tbl)

            cm_headers = ["", "2nd Test: Maybe", "2nd Test: Exclude", "Total"]
            for j, h in enumerate(cm_headers):
                set_cell(cm_tbl.cell(0, j), h, bold=True, font_size=Pt(8))
                shade(cm_tbl.cell(0, j), "D9E2F3")

            tp_, fp_, fn_, tn_ = r["tp"], r["fp"], r["fn"], r["tn"]
            cm_data = [
                ("1st Test: Maybe",   str(tp_), str(fp_), str(tp_ + fp_)),
                ("1st Test: Exclude", str(fn_), str(tn_), str(fn_ + tn_)),
                ("Total", str(tp_ + fn_), str(fp_ + tn_), str(r["n_total"])),
            ]
            for i, (label, *vals) in enumerate(cm_data, start=1):
                set_cell(cm_tbl.cell(i, 0), label, bold=True,
                         font_size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT)
                for j2, v in enumerate(vals, start=1):
                    set_cell(cm_tbl.cell(i, j2), v, font_size=Pt(8))
                if i == 1:
                    shade(cm_tbl.cell(i, 1), "D5F5E3")
                if i == 2:
                    shade(cm_tbl.cell(i, 2), "D5F5E3")

            doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  5. FALSE NEGATIVES (this project)
    # ==================================================================
    add_heading(doc, "5. False Negatives", level=1)
    add_note(doc, "Articles included by the human (maybe) but excluded by the AI (exclude). "
             "Most critical errors for systematic reviews.")

    if not proj_fn:
        p = doc.add_paragraph("No false negative data available for this project.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Negatives Count by Model", level=2)

        fn_headers = ["Model", "Test", "Total Paired", "False Neg.", "% of Total"]
        fn_rows = []
        for mn in sorted(proj_fn.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_fn[mn].keys()):
                r = proj_fn[mn][test_num]
                fn_count = r["fn"]
                n_total = r["n_paired"]
                fn_pct = fn_count / n_total * 100 if n_total > 0 else 0
                fn_rows.append([
                    model_name, f"{test_num}º",
                    str(n_total), str(fn_count), f"{fn_pct:.1f}%",
                ])

        tbl = doc.add_table(rows=1 + len(fn_rows), cols=len(fn_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fn_headers)
        for i, row_data in enumerate(fn_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 3:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        doc.add_paragraph()

    # ==================================================================
    #  6. FALSE POSITIVES (this project)
    # ==================================================================
    add_heading(doc, "6. False Positives", level=1)
    add_note(doc, "Articles excluded by the human (exclude) but included by the AI (maybe).")

    if not proj_fp:
        p = doc.add_paragraph("No false positive data available for this project.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Positives Count by Model", level=2)

        fp_headers = ["Model", "Test", "Total Paired", "False Pos.", "% of Total",
                      "Human Excl. Articles", "FP Rate (of excl.)"]
        fp_rows = []
        for mn in sorted(proj_fp.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_fp[mn].keys()):
                r = proj_fp[mn][test_num]
                fp_count = r["fp"]
                n_total = r["n_paired"]
                n_excl_human = r["fp"] + r["tn"]
                fp_pct = fp_count / n_total * 100 if n_total > 0 else 0
                fp_rate = fp_count / n_excl_human * 100 if n_excl_human > 0 else 0
                fp_rows.append([
                    model_name, f"{test_num}º",
                    str(n_total), str(fp_count), f"{fp_pct:.1f}%",
                    str(n_excl_human), f"{fp_rate:.1f}%",
                ])

        tbl = doc.add_table(rows=1 + len(fp_rows), cols=len(fp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fp_headers)
        for i, row_data in enumerate(fp_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  7. HOURS SAVED PER EXECUTION (this project)
    # ==================================================================
    if proj_diag:
        add_heading(doc, "7. Full-Text Hours Saved per Execution", level=1)
        add_note(doc, "Hours of full-text reading saved per AI run, based on this "
                 "project's data. Assumes 2 reviewers × 5 min per article. "
                 "Negative values mean the AI was more inclusive than the human.")

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Full-Text Hours Saved per Execution", level=2)

        hs_headers = ["Model", "Test", "Human TIAB Inc.",
                       "AI TIAB Inc.", "Articles Saved", "Hours Saved (FT)"]
        hs_rows = []
        for mn in sorted(proj["models"].keys()):
            model_info = proj["models"][mn]
            model_name = model_info["name"]
            for test_num in sorted(model_info["tests"].keys()):
                d = proj_diag.get(mn, {}).get(test_num)
                if not d:
                    continue
                ai_inc = d["tp"] + d["fp"]
                hu_inc = d["tp"] + d["fn"]
                arts_saved = hu_inc - ai_inc
                hrs_saved = arts_saved * 2 * 5 / 60.0
                hs_rows.append([
                    model_name, f"{test_num}o",
                    str(hu_inc), str(ai_inc),
                    f"{arts_saved:+d}",
                    f"{hrs_saved:+.1f}h",
                ])

        if hs_rows:
            tbl = doc.add_table(rows=1 + len(hs_rows), cols=len(hs_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(tbl)
            header_row(tbl, hs_headers)
            for i, row_data in enumerate(hs_rows):
                for j, val in enumerate(row_data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                    if j == 5:
                        try:
                            num = float(val.replace("h", "").replace("+", ""))
                            if num > 0:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif num < 0:
                                shade(tbl.cell(i + 1, j), "FFD6D6")
                        except ValueError:
                            pass

            doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  APPENDIX A — TIAB FALSE POSITIVES (this project)
    # ==================================================================
    add_heading(doc, "Appendix A. TIAB False Positives by Run", level=1)
    add_note(doc, "Articles included by the AI but excluded by the human screener at the TIAB stage.")

    has_fp = any(
        any(proj_fp.get(mn, {}).get(tn2, {}).get("fp_titles", [])
            for tn2 in proj_fp.get(mn, {}))
        for mn in proj_fp
    )

    if not has_fp:
        p = doc.add_paragraph("No false positive data for this project.")
        p.runs[0].font.italic = True
    else:
        all_fp_dict = {}
        run_keys_fp = []
        for mn in sorted(proj_fp.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_fp[mn].keys()):
                r = proj_fp[mn][test_num]
                label = f"{model_name} {test_num}º"
                run_keys_fp.append((mn, test_num, label))
                for art in r.get("fp_articles", []):
                    t = art.get("title", "")
                    if t and t not in all_fp_dict:
                        all_fp_dict[t] = art
                for t in r.get("fp_titles", []):
                    if t and t not in all_fp_dict:
                        all_fp_dict[t] = {"title": t, "abstract": ""}

        if not all_fp_dict:
            p = doc.add_paragraph("No false positives for this project.")
            p.runs[0].font.italic = True
        else:
            add_note(doc, f"{len(all_fp_dict)} unique article(s) classified as false positive.")

            for idx, title in enumerate(sorted(all_fp_dict.keys()), 1):
                art = all_fp_dict[title]
                abstract_text = art.get("abstract", "") or ""
                if pd.isna(abstract_text):
                    abstract_text = ""

                fp_models = []
                for mn, test_num, label in run_keys_fp:
                    r = proj_fp[mn][test_num]
                    if title in r.get("fp_titles", []):
                        fp_models.append(label)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.space_before = Pt(6)
                p_title.paragraph_format.space_after = Pt(2)
                run_num = p_title.add_run(f"{idx}. ")
                run_num.bold = True
                run_num.font.size = Pt(9)
                run_num.font.name = "Times New Roman"
                run_t = p_title.add_run(str(title))
                run_t.bold = True
                run_t.font.size = Pt(9)
                run_t.font.name = "Times New Roman"

                p_models = doc.add_paragraph()
                p_models.paragraph_format.left_indent = Pt(18)
                p_models.paragraph_format.space_before = Pt(0)
                p_models.paragraph_format.space_after = Pt(2)
                run_lbl = p_models.add_run("False positive in: ")
                run_lbl.bold = True
                run_lbl.font.size = Pt(8)
                run_lbl.font.name = "Times New Roman"
                run_lbl.font.color.rgb = RGBColor(200, 0, 0)
                run_mlist = p_models.add_run(", ".join(fp_models) if fp_models else "—")
                run_mlist.font.size = Pt(8)
                run_mlist.font.name = "Times New Roman"

                if abstract_text.strip():
                    p_abs = doc.add_paragraph()
                    p_abs.paragraph_format.left_indent = Pt(18)
                    p_abs.paragraph_format.space_after = Pt(6)
                    run_ab = p_abs.add_run(str(abstract_text))
                    run_ab.italic = True
                    run_ab.font.size = Pt(8)
                    run_ab.font.name = "Times New Roman"

    doc.add_page_break()

    # ==================================================================
    #  APPENDIX B — FULLTEXT MISSED ARTICLES (this project)
    # ==================================================================
    add_heading(doc, "Appendix B. Fulltext Missed Articles", level=1)
    add_note(doc, "Articles included in the human fulltext review but excluded by the "
             "AI in at least one run.")

    has_missed = any(
        any(proj_ft.get(mn, {}).get(tn2, {}).get("missed_articles", [])
            for tn2 in proj_ft.get(mn, {}))
        for mn in proj_ft
    )

    if not has_missed:
        p = doc.add_paragraph("No missed fulltext articles for this project.")
        p.runs[0].font.italic = True
    else:
        run_keys_miss = []
        for mn in sorted(proj_ft.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_ft[mn].keys()):
                label = f"{model_name} {test_num}º"
                run_keys_miss.append((mn, test_num, label))

        seen_titles = set()
        unique_missed = []
        for mn in sorted(proj_ft.keys()):
            for test_num in sorted(proj_ft[mn].keys()):
                r = proj_ft[mn][test_num]
                for art in r.get("missed_articles", []):
                    t = art.get("title", "")
                    if t and t not in seen_titles:
                        seen_titles.add(t)
                        unique_missed.append(art)

        if not unique_missed:
            p = doc.add_paragraph("No missed articles for this project.")
            p.runs[0].font.italic = True
        else:
            add_note(doc, f"{len(unique_missed)} unique article(s) missed by the AI.")

            # Matrix table
            tn_num = next_table()
            add_heading(doc, f"Table {tn_num}. Fulltext Articles Missed by AI", level=3)
            add_note(doc, "Each column shows a model/test run: ✗ = missed, ✓ = captured.")

            miss_headers = ["#", "Article Title"] + [rk[2] for rk in run_keys_miss]
            n_cols = len(miss_headers)

            miss_tbl = doc.add_table(rows=1, cols=n_cols)
            miss_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(miss_tbl)
            header_row(miss_tbl, miss_headers)

            sorted_missed = sorted(unique_missed, key=lambda a: a.get("title", ""))
            for idx, art in enumerate(sorted_missed, 1):
                title_app = art.get("title", "—") or "—"
                row = miss_tbl.add_row()
                set_cell(row.cells[0], str(idx), font_size=Pt(7))
                set_cell(row.cells[1], str(title_app)[:150], font_size=Pt(7),
                         align=WD_ALIGN_PARAGRAPH.LEFT)

                for col_idx, (mn, test_num, _label) in enumerate(run_keys_miss, start=2):
                    r = proj_ft[mn][test_num]
                    missed_in_run = r.get("missed_titles", [])
                    if title_app in missed_in_run:
                        set_cell(row.cells[col_idx], "✗", font_size=Pt(7),
                                 color=RGBColor(200, 0, 0))
                        shade(row.cells[col_idx], "FFD6D6")
                    else:
                        set_cell(row.cells[col_idx], "✓", font_size=Pt(7),
                                 color=RGBColor(0, 128, 0))
                        shade(row.cells[col_idx], "D5F5E3")

            doc.add_paragraph()

            # Detailed articles
            add_heading(doc, "Missed Articles Detail", level=3)

            for idx, art in enumerate(sorted_missed, 1):
                title_app = art.get("title", "—") or "—"
                abstract_app = art.get("abstract", "—") or "—"
                if pd.isna(abstract_app):
                    abstract_app = "—"

                missed_by = []
                captured_by = []
                for mn, test_num, label in run_keys_miss:
                    r = proj_ft[mn][test_num]
                    if title_app in r.get("missed_titles", []):
                        missed_by.append(label)
                    else:
                        captured_by.append(label)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.space_before = Pt(6)
                p_title.paragraph_format.space_after = Pt(2)
                run_num = p_title.add_run(f"{idx}. ")
                run_num.bold = True
                run_num.font.size = Pt(9)
                run_num.font.name = "Times New Roman"
                run_t = p_title.add_run(str(title_app))
                run_t.bold = True
                run_t.font.size = Pt(9)
                run_t.font.name = "Times New Roman"

                p_models = doc.add_paragraph()
                p_models.paragraph_format.left_indent = Pt(18)
                p_models.paragraph_format.space_before = Pt(0)
                p_models.paragraph_format.space_after = Pt(2)
                run_lbl = p_models.add_run("Missed by: ")
                run_lbl.bold = True
                run_lbl.font.size = Pt(8)
                run_lbl.font.name = "Times New Roman"
                run_lbl.font.color.rgb = RGBColor(200, 0, 0)
                run_mlist = p_models.add_run(", ".join(missed_by) if missed_by else "—")
                run_mlist.font.size = Pt(8)
                run_mlist.font.name = "Times New Roman"
                if captured_by:
                    run_cap = p_models.add_run("  |  Captured by: ")
                    run_cap.bold = True
                    run_cap.font.size = Pt(8)
                    run_cap.font.name = "Times New Roman"
                    run_cap.font.color.rgb = RGBColor(0, 128, 0)
                    run_clist = p_models.add_run(", ".join(captured_by))
                    run_clist.font.size = Pt(8)
                    run_clist.font.name = "Times New Roman"

                p_abs = doc.add_paragraph()
                p_abs.paragraph_format.left_indent = Pt(18)
                p_abs.paragraph_format.space_after = Pt(6)
                run_ab = p_abs.add_run(str(abstract_app))
                run_ab.italic = True
                run_ab.font.size = Pt(8)
                run_ab.font.name = "Times New Roman"

    return doc
