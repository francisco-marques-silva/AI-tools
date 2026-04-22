"""
report_generator.py — Build the unified Word (.docx) report.

Contains only the generate_report() function and its local helpers.
No chart generation — charts are produced separately by graphic.py / graphic.R
from the data_grafics_*.xlsx file exported by chart_data.py.
"""

import datetime
from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .utils import fmt, fmt_pct, normalise_model_name, load_file
from .docx_helpers import (
    shade, set_cell, add_borders,
    add_heading, add_note, header_row,
)


# =====================================================================
#  generate_report
# =====================================================================

def generate_report(projects, metadados, all_results, output_dir: Path):
    """Generate the unified Word document with all analyses (tables only)."""

    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    table_counter = [0]

    def next_table():
        table_counter[0] += 1
        return table_counter[0]

    # ==================================================================
    #  COVER / TITLE
    # ==================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Unified AI Screening Analysis Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated on: {ts}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    proj_names = [projects[pn]["name"] for pn in sorted(projects.keys())]
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Projects analyzed: {', '.join(proj_names)}")
    run.font.size = Pt(10)

    doc.add_page_break()

    # ==================================================================
    #  METHODOLOGICAL NOTES & REPORT GUIDE
    # ==================================================================
    add_heading(doc, "Methodological Notes & Report Guide", level=1)

    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(8)
    run = intro_p.add_run(
        "This report consolidates all analyses comparing AI-based screening decisions "
        "against human reviewer decisions across multiple projects and models. "
        "The sections below are organized progressively: from data validation through "
        "diagnostic performance, reproducibility, error analysis, and cost-efficiency."
    )
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    section_guide = [
        ("Section 1 — Data Validation",
         "Inventories all detected input files (AI results, human TIAB, Fulltext, Listfinal, metadata) "
         "and checks correspondence between AI result files and metadata entries."),
        ("Section 2 — Metadata and Costs",
         "Displays the full execution metadata table (model, parameters, execution time, tokens, cost) "
         "and provides cost summaries grouped by project and by model."),
        ("Section 3 — Diagnostic Analysis (AI vs Human TIAB)",
         "For each project × model × test, computes sensitivity, specificity, PPV, NPV, accuracy, F1, "
         "and Cohen's Kappa by comparing AI screening decisions against the human TIAB reference."),
        ("Section 4 — Fulltext Verification",
         "Measures the capture rate: the proportion of articles selected for full-text reading "
         "that the AI would have retained during TIAB screening."),
        ("Section 5 — Listfinal Verification (Gold Standard)",
         "The definitive performance measure. Shows the capture rate over the final included articles "
         "(post full-text reading)."),
        ("Section 6 — Test-Retest (Reproducibility)",
         "Compares two independent runs of the same model on the same dataset, measuring exact agreement, "
         "binarized agreement, and Cohen's Kappa with 95% CI."),
        ("Section 7 — False Negatives",
         "Lists articles included by the human reviewer but excluded by the AI."),
        ("Section 8 — False Positives",
         "Lists articles excluded by the human reviewer but included by the AI."),
        ("Section 9 — General Comparative Table",
         "A consolidated table presenting all key metrics for every model × project × test."),
        ("Section 10 — Cost-Effectiveness",
         "Relates cost (USD) to mean sensitivity per model."),
        ("Section 11 — Workload Reduction",
         "Compares human screening time against AI screening time per execution and per project."),
        ("Section 12 — Absolute Efficiency",
         "Combines selectivity and capture into an efficiency score. "
         "Efficiency Score = Listfinal Capture Rate × (1 − AI Positive Rate)."),
        ("Appendix A — TIAB False Positives",
         "Detailed per-article list with title, abstract, and which models flagged each article."),
        ("Appendix B — Fulltext Missed Articles",
         "Per-article matrix showing which models missed each fulltext-selected article."),
    ]
    for title_text, desc in section_guide:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run = p.add_run(desc)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    doc.add_paragraph()
    meth_title = doc.add_paragraph()
    run = meth_title.add_run("Key Methodological Definitions")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    method_notes = [
        ("Binarization", "AI decisions (include, maybe, exclude) are binarized: "
         "include and maybe → positive (passes screening), exclude → negative."),
        ("Gold Standard Hierarchy", "Three reference levels: (1) Human TIAB decisions — "
         "primary comparison; (2) Fulltext selection; (3) Listfinal — definitive measure."),
        ("Cohen's Kappa", "Interpretation follows Landis & Koch (1977): < 0 Poor; "
         "0.00–0.20 Slight; 0.21–0.40 Fair; 0.41–0.60 Moderate; "
         "0.61–0.80 Substantial; 0.81–1.00 Almost Perfect."),
        ("Workload Reduction", "Human screening time (time_human) compared with AI screening "
         "time (time_ia). Speed factor = human time / AI time."),
        ("Absolute Efficiency", "Efficiency Score = LF Capture Rate × (1 − AI Positive Rate)."),
        ("Cost-Effectiveness", "Cost per sensitivity point = average cost / (sensitivity × 100)."),
    ]
    for title_text, text in method_notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    doc.add_page_break()

    # ==================================================================
    #  SECTION 1 — DATA VALIDATION
    # ==================================================================
    add_heading(doc, "1. Data Validation", level=1)

    tn = next_table()
    add_heading(doc, f"Table {tn}. Detected Files Inventory", level=2)

    inventory_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        n_models = len(proj["models"])
        n_ai_files = sum(len(m["tests"]) for m in proj["models"].values())
        has_tiab = "Yes" if proj["human_tiab"] else "No"
        has_ft = "Yes" if proj["human_fulltext"] else "No"
        has_lf = "Yes" if proj.get("human_listfinal") else "No"
        lf_n = ""
        if proj.get("human_listfinal"):
            try:
                lf_df = load_file(str(proj["human_listfinal"]))
                lf_n = str(len(lf_df))
            except Exception:
                lf_n = "?"
        inventory_rows.append([
            proj["name"], str(n_models), str(n_ai_files), has_tiab, has_ft, has_lf, lf_n
        ])

    inv_headers = ["Project", "Models", "AI Files", "Human TIAB", "Human Fulltext",
                   "Listfinal", "Listfinal N"]
    tbl = doc.add_table(rows=1 + len(inventory_rows), cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, inv_headers)
    for i, row_data in enumerate(inventory_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
            if j == 3 and val == "No":
                shade(tbl.cell(i + 1, j), "FFD6D6")
            if j in (4, 5) and val == "No":
                shade(tbl.cell(i + 1, j), "FFF3CD")

    doc.add_paragraph()

    validation_issues = all_results.get("validation_issues", [])
    if validation_issues:
        add_heading(doc, "Validation Notes", level=3)
        for issue in validation_issues:
            p = doc.add_paragraph(f"• {issue}")
            p.runs[0].font.size = Pt(9)
    else:
        p = doc.add_paragraph("✓ All files were validated successfully.")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()

    # ==================================================================
    #  SECTION 2 — METADATA AND COSTS
    # ==================================================================
    if metadados is not None and not metadados.empty:
        add_heading(doc, "2. Metadata and Costs", level=1)

        tn = next_table()
        add_heading(doc, f"Table {tn}. Execution Metadata", level=2)
        add_note(doc, "Details of each execution: model, parameters, time, tokens, and costs (USD).")

        meta_cols = ["Project", "Code", "Model", "Parameters", "Version",
                     "Time", "Tokens In", "Tokens Out", "Cost ($)"]
        meta_rows = []
        for _, row in metadados.iterrows():
            meta_rows.append([
                str(row.get("project", "")),
                str(row.get("code", "")),
                str(row.get("model", "")),
                str(row.get("parameter", "")),
                str(row.get("version", "")),
                str(row.get("time_ia", "")),
                str(int(row["tokens input"])) if pd.notna(row.get("tokens input")) else "-",
                str(int(row["tokens_output"])) if pd.notna(row.get("tokens_output")) else "-",
                f"{row['cost_total']:.2f}" if pd.notna(row.get("cost_total")) else "-",
            ])

        tbl = doc.add_table(rows=1 + len(meta_rows), cols=len(meta_cols))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, meta_cols)
        for i, row_data in enumerate(meta_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)

        doc.add_paragraph()

        # Cost Summary by Project
        tn = next_table()
        add_heading(doc, f"Table {tn}. Cost Summary by Project", level=2)

        meta_valid = metadados.dropna(subset=["project"])
        cost_summary = meta_valid.groupby("project").agg(
            n_execucoes=("cost_total", "count"),
            custo_total=("cost_total", "sum"),
            custo_medio=("cost_total", "mean"),
            tokens_in_total=("tokens input", "sum"),
            tokens_out_total=("tokens_output", "sum"),
        ).reset_index()

        cost_headers = ["Project", "Executions", "Total Cost ($)", "Avg Cost ($)",
                        "Tokens In (total)", "Tokens Out (total)"]
        tbl = doc.add_table(rows=1 + len(cost_summary), cols=6)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cost_headers)
        for i, (_, row) in enumerate(cost_summary.iterrows()):
            vals = [
                str(row["project"]),
                str(int(row["n_execucoes"])),
                f"{row['custo_total']:.2f}",
                f"{row['custo_medio']:.2f}",
                str(int(row["tokens_in_total"])),
                str(int(row["tokens_out_total"])),
            ]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

        # Average Cost per Model
        tn = next_table()
        add_heading(doc, f"Table {tn}. Average Cost per Model", level=2)

        cost_model = metadados.dropna(subset=["project"]).copy()
        cost_model["model_norm"] = cost_model["model"].apply(normalise_model_name)
        cost_by_model = cost_model.groupby("model").agg(
            n=("cost_total", "count"),
            custo_medio=("cost_total", "mean"),
            custo_total=("cost_total", "sum"),
        ).reset_index().sort_values("custo_medio")

        cmod_headers = ["Model", "Executions", "Avg Cost ($)", "Total Cost ($)"]
        tbl = doc.add_table(rows=1 + len(cost_by_model), cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cmod_headers)
        for i, (_, row) in enumerate(cost_by_model.iterrows()):
            vals = [str(row["model"]), str(int(row["n"])),
                    f"{row['custo_medio']:.2f}", f"{row['custo_total']:.2f}"]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 3 — DIAGNOSTIC ANALYSIS
    # ==================================================================
    section_num = 3
    add_heading(doc, f"{section_num}. TIAB Agreement Analysis (AI vs Human Screener)", level=1)
    add_note(doc, "Comparison between AI and human screener decisions at the TIAB level. "
             "Note: the human TIAB decision is an intermediate reference, not the final gold standard.")

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        proj_diag = all_results.get("diagnostic", {}).get(pn, {})

        if not proj_diag:
            p = doc.add_paragraph(f"Project {proj['name']}: no human reference (TIAB).")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True
            continue

        add_heading(doc, f"Project: {proj['name']}", level=2)

        # Metric comparison table
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Model Comparison — {proj['name']}", level=3)
        add_note(doc, "Diagnostic metrics for each model and test, compared to the human gold standard.")

        comp_headers = ["Model", "Test", "N", "TP", "FP", "FN", "TN",
                        "Sens.", "Spec.", "PPV", "NPV", "Acc.", "F1", "Kappa", "95% CI", "Interpretation"]
        comp_rows = []

        for mn in sorted(proj_diag.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_diag[mn].keys()):
                r = proj_diag[mn][test_num]
                if r is None:
                    continue
                m = r["metrics"]
                comp_rows.append([
                    model_name,
                    f"{test_num}º",
                    str(r["n_paired"]),
                    str(r["tp"]), str(r["fp"]), str(r["fn"]), str(r["tn"]),
                    fmt_pct(m["Sensitivity"]),
                    fmt_pct(m["Specificity"]),
                    fmt_pct(m["PPV (Precision)"]),
                    fmt_pct(m["NPV"]),
                    fmt_pct(m["Accuracy"]),
                    fmt(m["F1 Score"], 3),
                    fmt(r["kappa"], 3),
                    f"[{fmt(r['kappa_ci_lo'], 2)}, {fmt(r['kappa_ci_hi'], 2)}]",
                    r["kappa_interp"],
                ])

        tbl = doc.add_table(rows=1 + len(comp_rows), cols=len(comp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, comp_headers)
        for i, row_data in enumerate(comp_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
                if j == 7 and val not in ("-", "N/A"):
                    try:
                        v = float(val.replace("%", ""))
                        if v >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif v < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass

        doc.add_paragraph()

        # Confusion matrices
        for mn in sorted(proj_diag.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_diag[mn].keys()):
                r = proj_diag[mn][test_num]
                if r is None:
                    continue

                tn_num = next_table()
                add_heading(doc, f"Table {tn_num}. Confusion Matrix — {model_name} (Test {test_num})", level=3)

                cm_tbl = doc.add_table(rows=4, cols=4)
                cm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                add_borders(cm_tbl)

                cm_headers = ["", "Human: Maybe", "Human: Exclude", "Total"]
                for j, h in enumerate(cm_headers):
                    set_cell(cm_tbl.cell(0, j), h, bold=True, font_size=Pt(8))
                    shade(cm_tbl.cell(0, j), "D9E2F3")

                cm_data = [
                    ("AI: Maybe",   str(r["tp"]), str(r["fp"]), str(r["tp"] + r["fp"])),
                    ("AI: Exclude", str(r["fn"]), str(r["tn"]), str(r["fn"] + r["tn"])),
                    ("Total",       str(r["tp"] + r["fn"]), str(r["fp"] + r["tn"]),
                     str(r["tp"] + r["fp"] + r["fn"] + r["tn"])),
                ]
                for i, (label, *vals) in enumerate(cm_data, start=1):
                    set_cell(cm_tbl.cell(i, 0), label, bold=True,
                             font_size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT)
                    for j2, v in enumerate(vals, start=1):
                        set_cell(cm_tbl.cell(i, j2), v, font_size=Pt(8))
                    if i == 1:
                        shade(cm_tbl.cell(i, 1), "D5F5E3")  # TP
                    if i == 2:
                        shade(cm_tbl.cell(i, 2), "D5F5E3")  # TN

                doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 4 — FULLTEXT VERIFICATION
    # ==================================================================
    section_num = 4
    add_heading(doc, f"{section_num}. Fulltext Verification (Capture Rate)", level=1)
    add_note(doc, "Checks whether articles included in the final review (fulltext) would have "
             "been retained by the AI during TIAB screening.")

    ft_results = all_results.get("fulltext", {})
    if not ft_results:
        p = doc.add_paragraph("No project with fulltext data available.")
        p.runs[0].font.italic = True
    else:
        for pn in sorted(ft_results.keys()):
            proj = projects[pn]
            add_heading(doc, f"Project: {proj['name']}", level=2)

            proj_ft = ft_results[pn]

            tn_num = next_table()
            add_heading(doc, f"Table {tn_num}. Fulltext Capture Rate — {proj['name']}", level=3)

            ft_headers = ["Model", "Test", "FT Articles", "Found",
                          "Captured", "Missed", "Capture Rate", "Miss Rate"]
            ft_rows = []
            for mn in sorted(proj_ft.keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(proj_ft[mn].keys()):
                    r = proj_ft[mn][test_num]
                    ft_rows.append([
                        model_name, f"{test_num}º",
                        str(r["n_fulltext"]), str(r["n_found"]),
                        str(r["n_captured"]), str(r["n_missed"]),
                        fmt_pct(r["capture_rate"]), fmt_pct(r["miss_rate"]),
                    ])

            tbl = doc.add_table(rows=1 + len(ft_rows), cols=len(ft_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(tbl)
            header_row(tbl, ft_headers)
            for i, row_data in enumerate(ft_rows):
                for j, val in enumerate(row_data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                    if j == 6 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val >= 95:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val < 80:
                                shade(tbl.cell(i + 1, j), "FFD6D6")
                        except ValueError:
                            pass
                    if j == 5:
                        try:
                            if int(val) > 0:
                                shade(tbl.cell(i + 1, j), "FFF3CD")
                        except ValueError:
                            pass

            doc.add_paragraph()

            all_missed = set()
            for mn in sorted(proj_ft.keys()):
                for test_num in sorted(proj_ft[mn].keys()):
                    r = proj_ft[mn][test_num]
                    for t in r.get("missed_titles", []):
                        all_missed.add(t)
            if all_missed:
                p = doc.add_paragraph(
                    f"Note: {len(all_missed)} unique article(s) were missed by at least one AI run. "
                    "See Appendix (Fulltext Missed Articles) for details."
                )
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.italic = True

            doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 5 — LISTFINAL VERIFICATION
    # ==================================================================
    section_num = 5
    add_heading(doc, f"{section_num}. Listfinal Verification (True Gold Standard Capture Rate)", level=1)
    add_note(doc, "Checks whether articles in the final included list (Listfinal) "
             "would have been retained by the AI during TIAB screening.")

    lf_results = all_results.get("listfinal", {})
    if not lf_results:
        p = doc.add_paragraph("No project with Listfinal data available.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Listfinal Capture Rate — All Projects", level=2)
        add_note(doc, "How many of the final included articles would have been retained by each AI model.")

        lf_headers = ["Project", "Model", "Test", "Listfinal N", "Found",
                       "Captured", "Missed", "Capture Rate", "Miss Rate"]
        lf_rows = []
        for pn in sorted(lf_results.keys()):
            proj = projects[pn]
            for mn in sorted(lf_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(lf_results[pn][mn].keys()):
                    r = lf_results[pn][mn][test_num]
                    lf_rows.append([
                        proj["name"], model_name, f"{test_num}o",
                        str(r["n_listfinal"]), str(r["n_found"]),
                        str(r["n_captured"]), str(r["n_missed"]),
                        fmt_pct(r["capture_rate"]), fmt_pct(r["miss_rate"]),
                    ])

        tbl = doc.add_table(rows=1 + len(lf_rows), cols=len(lf_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, lf_headers)
        for i, row_data in enumerate(lf_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 7 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass
                if j == 6:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFF3CD")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 6 — TEST-RETEST
    # ==================================================================
    section_num = 6
    add_heading(doc, f"{section_num}. Test-Retest (Reproducibility)", level=1)
    add_note(doc, "Compares two runs of the same model on the same dataset.")

    tr_results = all_results.get("test_retest", {})
    if not tr_results:
        p = doc.add_paragraph("No test-retest pairs found.")
        p.runs[0].font.italic = True
    else:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Test-Retest Summary — All Projects", level=2)

        tr_headers = ["Project", "Model", "N", "Exact Agree.",
                      "Binary Agree.", "Discordant", "Kappa", "95% CI", "Interpretation"]
        tr_rows = []
        for pn in sorted(tr_results.keys()):
            proj = projects[pn]
            for mn in sorted(tr_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                r = tr_results[pn][mn]
                tr_rows.append([
                    proj["name"], model_name,
                    str(r["n_total"]),
                    f"{r['exact_match']} ({fmt_pct(r['exact_pct'])})",
                    f"{r['binary_match']} ({fmt_pct(r['binary_pct'])})",
                    str(r["n_discordant"]),
                    fmt(r["kappa"], 3),
                    f"[{fmt(r['kappa_ci_lo'], 2)}, {fmt(r['kappa_ci_hi'], 2)}]",
                    r["kappa_interp"],
                ])

        tbl = doc.add_table(rows=1 + len(tr_rows), cols=len(tr_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, tr_headers)
        for i, row_data in enumerate(tr_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)

        doc.add_paragraph()

        # Detail by project
        for pn in sorted(tr_results.keys()):
            proj = projects[pn]
            add_heading(doc, f"Project: {proj['name']}", level=2)

            for mn in sorted(tr_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                r = tr_results[pn][mn]

                tn_num = next_table()
                add_heading(doc, f"Table {tn_num}. Binary Test-Retest Matrix — {model_name} ({proj['name']})", level=3)

                cm_tbl = doc.add_table(rows=4, cols=4)
                cm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                add_borders(cm_tbl)

                cm_headers = ["", "2nd Test: Maybe", "2nd Test: Exclude", "Total"]
                for j, h in enumerate(cm_headers):
                    set_cell(cm_tbl.cell(0, j), h, bold=True, font_size=Pt(8))
                    shade(cm_tbl.cell(0, j), "D9E2F3")

                tp_, fp_, fn_, tn_ = r["tp"], r["fp"], r["fn"], r["tn"]
                cm_data = [
                    ("1st Test: Maybe",   str(tp_), str(fp_), str(tp_ + fp_)),
                    ("1st Test: Exclude", str(fn_), str(tn_), str(fn_ + tn_)),
                    ("Total", str(tp_ + fn_), str(fp_ + tn_), str(r["n_total"])),
                ]
                for i, (label, *vals) in enumerate(cm_data, start=1):
                    set_cell(cm_tbl.cell(i, 0), label, bold=True,
                             font_size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT)
                    for j2, v in enumerate(vals, start=1):
                        set_cell(cm_tbl.cell(i, j2), v, font_size=Pt(8))
                    if i == 1:
                        shade(cm_tbl.cell(i, 1), "D5F5E3")
                    if i == 2:
                        shade(cm_tbl.cell(i, 2), "D5F5E3")

                doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 7 — FALSE NEGATIVES
    # ==================================================================
    section_num = 7
    add_heading(doc, f"{section_num}. False Negatives Analysis", level=1)
    add_note(doc, "Articles included by the human (maybe) but excluded by the AI (exclude). "
             "False negatives are the most critical in systematic review screening.")

    fn_results = all_results.get("false_negatives", {})
    if fn_results:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Negatives Count by Model", level=2)

        fn_headers = ["Project", "Model", "Test", "Total Paired", "False Neg.", "% of Total"]
        fn_rows = []
        for pn in sorted(fn_results.keys()):
            proj = projects[pn]
            for mn in sorted(fn_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(fn_results[pn][mn].keys()):
                    r = fn_results[pn][mn][test_num]
                    fn_count = r["fn"]
                    n_total = r["n_paired"]
                    fn_pct = fn_count / n_total * 100 if n_total > 0 else 0
                    fn_rows.append([
                        proj["name"], model_name, f"{test_num}º",
                        str(n_total), str(fn_count), f"{fn_pct:.1f}%",
                    ])

        tbl = doc.add_table(rows=1 + len(fn_rows), cols=len(fn_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fn_headers)
        for i, row_data in enumerate(fn_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 4:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        doc.add_paragraph()
    else:
        p = doc.add_paragraph("No false negative data (requires human reference).")
        p.runs[0].font.italic = True

    # ==================================================================
    #  SECTION 8 — FALSE POSITIVES
    # ==================================================================
    section_num = 8
    add_heading(doc, f"{section_num}. False Positives Analysis", level=1)
    add_note(doc, "Articles excluded by the human (exclude) but included by the AI (maybe).")

    fp_results = all_results.get("false_positives", {})
    if fp_results:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Positives Count by Model", level=2)

        fp_headers = ["Project", "Model", "Test", "Total Paired", "False Pos.", "% of Total",
                      "Human Excl. Articles", "FP Rate (of excl.)"]
        fp_rows = []
        for pn in sorted(fp_results.keys()):
            proj = projects[pn]
            for mn in sorted(fp_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(fp_results[pn][mn].keys()):
                    r = fp_results[pn][mn][test_num]
                    fp_count = r["fp"]
                    n_total = r["n_paired"]
                    n_excl_human = r["fp"] + r["tn"]
                    fp_pct = fp_count / n_total * 100 if n_total > 0 else 0
                    fp_rate = fp_count / n_excl_human * 100 if n_excl_human > 0 else 0
                    fp_rows.append([
                        proj["name"], model_name, f"{test_num}º",
                        str(n_total), str(fp_count), f"{fp_pct:.1f}%",
                        str(n_excl_human), f"{fp_rate:.1f}%",
                    ])

        tbl = doc.add_table(rows=1 + len(fp_rows), cols=len(fp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fp_headers)
        for i, row_data in enumerate(fp_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()
    else:
        p = doc.add_paragraph("No false positive data (requires human reference).")
        p.runs[0].font.italic = True

    doc.add_page_break()

    # ==================================================================
    #  SECTION 9 — GENERAL COMPARATIVE TABLE
    # ==================================================================
    section_num = 9
    add_heading(doc, f"{section_num}. General Comparative Table", level=1)
    add_note(doc, "Consolidated view of all metrics by project × model × test.")

    diag = all_results.get("diagnostic", {})
    ft = all_results.get("fulltext", {})
    lf = all_results.get("listfinal", {})
    tr = all_results.get("test_retest", {})

    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. General Comparison — Diagnostic Performance and Reproducibility", level=2)

    big_headers = ["Project", "Model", "Test", "Sens.", "Spec.", "F1",
                   "Kappa (diag)", "FT Capture", "LF Capture", "Kappa (T-R)", "Cost ($)"]
    big_rows = []

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            model_info = proj["models"][mn]
            model_name = model_info["name"]
            for test_num in sorted(model_info["tests"].keys()):
                row_vals = [proj["name"], model_name, f"{test_num}o"]

                d = diag.get(pn, {}).get(mn, {}).get(test_num)
                if d:
                    row_vals.append(fmt_pct(d["metrics"]["Sensitivity"]))
                    row_vals.append(fmt_pct(d["metrics"]["Specificity"]))
                    row_vals.append(fmt(d["metrics"]["F1 Score"], 3))
                    row_vals.append(fmt(d["kappa"], 3))
                else:
                    row_vals.extend(["-", "-", "-", "-"])

                f_res = ft.get(pn, {}).get(mn, {}).get(test_num)
                row_vals.append(fmt_pct(f_res["capture_rate"]) if f_res else "-")

                lf_res = lf.get(pn, {}).get(mn, {}).get(test_num)
                row_vals.append(fmt_pct(lf_res["capture_rate"]) if lf_res else "-")

                tr_res = tr.get(pn, {}).get(mn)
                if tr_res and test_num == 1:
                    row_vals.append(fmt(tr_res["kappa"], 3))
                elif test_num == 2 and tr_res:
                    row_vals.append("↑")
                else:
                    row_vals.append("-")

                code = model_info["tests"][test_num]["code"]
                cost_val = "-"
                if metadados is not None:
                    meta_match = metadados[metadados["code"].astype(str) == str(code)]
                    if not meta_match.empty:
                        cost_val = f"{meta_match.iloc[0]['cost_total']:.2f}"
                row_vals.append(cost_val)

                big_rows.append(row_vals)

    tbl = doc.add_table(rows=1 + len(big_rows), cols=len(big_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, big_headers)
    for i, row_data in enumerate(big_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)

    doc.add_paragraph()

    # ==================================================================
    #  SECTION 10 — COST-EFFECTIVENESS
    # ==================================================================
    if metadados is not None and diag:
        section_num = 10
        add_heading(doc, f"{section_num}. Cost-Effectiveness Analysis", level=1)
        add_note(doc, "Relationship between cost (USD) and diagnostic performance.")

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Cost vs. Sensitivity per Model (test average)", level=2)

        cost_eff_headers = ["Model", "Project", "Avg Sens.", "Avg Spec.",
                            "Avg F1", "Avg Cost ($)", "Cost per Sens. point"]
        cost_eff_rows = []

        for pn in sorted(diag.keys()):
            proj = projects[pn]
            for mn in sorted(diag[pn].keys()):
                model_name = proj["models"][mn]["name"]
                tests = diag[pn][mn]
                sens_vals, spec_vals, f1_vals, cost_vals = [], [], [], []
                for tn2, r in tests.items():
                    if r is None:
                        continue
                    s = r["metrics"]["Sensitivity"]
                    if not np.isnan(s):
                        sens_vals.append(s)
                    sp = r["metrics"]["Specificity"]
                    if not np.isnan(sp):
                        spec_vals.append(sp)
                    f1v = r["metrics"]["F1 Score"]
                    if not np.isnan(f1v):
                        f1_vals.append(f1v)
                    code = proj["models"][mn]["tests"][tn2]["code"]
                    if metadados is not None:
                        meta_m = metadados[metadados["code"].astype(str) == str(code)]
                        if not meta_m.empty and pd.notna(meta_m.iloc[0]["cost_total"]):
                            cost_vals.append(meta_m.iloc[0]["cost_total"])

                avg_sens = np.mean(sens_vals) if sens_vals else float("nan")
                avg_spec = np.mean(spec_vals) if spec_vals else float("nan")
                avg_f1 = np.mean(f1_vals) if f1_vals else float("nan")
                avg_cost = np.mean(cost_vals) if cost_vals else float("nan")
                cost_per_sens = avg_cost / (avg_sens * 100) if (
                    not np.isnan(avg_cost) and not np.isnan(avg_sens) and avg_sens > 0
                ) else float("nan")

                cost_eff_rows.append([
                    model_name, proj["name"],
                    fmt_pct(avg_sens), fmt_pct(avg_spec),
                    fmt(avg_f1, 3),
                    fmt(avg_cost, 2) if not np.isnan(avg_cost) else "-",
                    fmt(cost_per_sens, 3) if not np.isnan(cost_per_sens) else "-",
                ])

        tbl = doc.add_table(rows=1 + len(cost_eff_rows), cols=len(cost_eff_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cost_eff_headers)
        for i, row_data in enumerate(cost_eff_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()

    # ==================================================================
    #  SECTION 11 — WORKLOAD REDUCTION
    # ==================================================================
    if metadados is not None:
        section_num = 11
        add_heading(doc, f"{section_num}. Workload Reduction Analysis", level=1)
        add_note(doc, "Compares human screening time (time_human) with AI screening time (time_ia).")

        def parse_td(val):
            if pd.isna(val):
                return float("nan")
            if isinstance(val, pd.Timedelta):
                return val.total_seconds() / 3600.0
            s = str(val).strip()
            if not s:
                return float("nan")
            try:
                return pd.to_timedelta(s).total_seconds() / 3600.0
            except Exception:
                return float("nan")

        def fmt_hours(h):
            if np.isnan(h):
                return "-"
            hr = int(h)
            mn = int((h - hr) * 60)
            return f"{hr}h {mn:02d}m" if hr > 0 else f"{mn}m"

        meta_work = metadados.copy()
        meta_work["_h_human"] = meta_work["time_human"].apply(parse_td)
        meta_work["_h_ia"] = meta_work["time_ia"].apply(parse_td)

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Workload Reduction — Per Execution", level=2)

        wr_headers = ["Project", "Model", "Code", "Human Time", "AI Time",
                      "Time Saved", "Reduction (%)", "Speed Factor"]
        wr_rows = []
        for _, row in meta_work.iterrows():
            h_human = row["_h_human"]
            h_ia = row["_h_ia"]
            saved = h_human - h_ia if not (np.isnan(h_human) or np.isnan(h_ia)) else float("nan")
            reduction = (saved / h_human * 100) if (not np.isnan(saved) and h_human > 0) else float("nan")
            factor = (h_human / h_ia) if (not np.isnan(h_human) and not np.isnan(h_ia) and h_ia > 0) else float("nan")
            wr_rows.append([
                str(row.get("project", "")),
                str(row.get("model", "")),
                str(row.get("code", "")),
                fmt_hours(h_human),
                fmt_hours(h_ia),
                fmt_hours(saved),
                f"{reduction:.1f}%" if not np.isnan(reduction) else "-",
                f"{factor:.0f}x" if not np.isnan(factor) else "-",
            ])

        tbl = doc.add_table(rows=1 + len(wr_rows), cols=len(wr_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, wr_headers)
        for i, row_data in enumerate(wr_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                if j == 6 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 90:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val >= 50:
                            shade(tbl.cell(i + 1, j), "E8F5E9")
                    except ValueError:
                        pass

        doc.add_paragraph()

        # Summary by project
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Workload Reduction Summary by Project", level=2)

        wr_proj = meta_work.groupby("project").agg(
            human_hours=("_h_human", "first"),
            avg_ia_hours=("_h_ia", "mean"),
            min_ia_hours=("_h_ia", "min"),
            max_ia_hours=("_h_ia", "max"),
            n_runs=("_h_ia", "count"),
        ).reset_index()

        wp_headers = ["Project", "Human Time", "Avg AI Time", "Fastest AI",
                      "Avg Reduction (%)", "Avg Speed Factor"]
        tbl = doc.add_table(rows=1 + len(wr_proj), cols=len(wp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, wp_headers)
        for i, (_, row) in enumerate(wr_proj.iterrows()):
            h_h = row["human_hours"]
            avg_ia = row["avg_ia_hours"]
            fast = row["min_ia_hours"]
            avg_red = ((h_h - avg_ia) / h_h * 100) if (not np.isnan(h_h) and not np.isnan(avg_ia) and h_h > 0) else float("nan")
            avg_fac = (h_h / avg_ia) if (not np.isnan(h_h) and not np.isnan(avg_ia) and avg_ia > 0) else float("nan")
            vals = [
                str(row["project"]),
                fmt_hours(h_h),
                fmt_hours(avg_ia),
                fmt_hours(fast),
                f"{avg_red:.1f}%" if not np.isnan(avg_red) else "-",
                f"{avg_fac:.0f}x" if not np.isnan(avg_fac) else "-",
            ]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

    # ==================================================================
    #  SECTION 12 — ABSOLUTE EFFICIENCY
    # ==================================================================
    diag_abs = all_results.get("diagnostic", {})
    lf_abs = all_results.get("listfinal", {})
    if diag_abs and lf_abs:
        section_num = 12
        add_heading(doc, f"{section_num}. Absolute Efficiency Analysis (TIAB Reduction vs Listfinal Retention)", level=1)
        add_note(doc, "Measures how much the AI reduces the TIAB workload while still retaining truly relevant articles.")

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Absolute Efficiency — TIAB Selection Volume vs Listfinal Retention", level=2)

        eff_headers = ["Project", "Model", "Test", "TIAB N",
                       "AI Positives", "AI Pos. (%)", "Human Positives", "Human Pos. (%)",
                       "Reduction", "LF Capture", "Efficiency Score"]
        eff_rows = []

        for pn in sorted(projects.keys()):
            proj = projects[pn]
            proj_diag = diag_abs.get(pn, {})
            proj_lf = lf_abs.get(pn, {})

            for mn in sorted(proj["models"].keys()):
                model_info = proj["models"][mn]
                model_name = model_info["name"]
                for test_num in sorted(model_info["tests"].keys()):
                    d = proj_diag.get(mn, {}).get(test_num)
                    l = proj_lf.get(mn, {}).get(test_num)
                    if not d:
                        continue

                    n_total = d["n_paired"]
                    ai_pos = d["tp"] + d["fp"]
                    hu_pos = d["tp"] + d["fn"]
                    ai_pct = ai_pos / n_total * 100 if n_total > 0 else float("nan")
                    hu_pct = hu_pos / n_total * 100 if n_total > 0 else float("nan")
                    reduction = ((hu_pos - ai_pos) / hu_pos * 100) if hu_pos > 0 else float("nan")

                    lf_cap = "-"
                    eff_score = "-"
                    if l:
                        cap_rate = l["capture_rate"]
                        lf_cap = fmt_pct(cap_rate)
                        if not np.isnan(cap_rate) and not np.isnan(ai_pct):
                            score = cap_rate * (1 - ai_pct / 100)
                            eff_score = fmt(score, 3)

                    eff_rows.append([
                        proj["name"], model_name, f"{test_num}o",
                        str(n_total),
                        str(ai_pos), f"{ai_pct:.1f}%" if not np.isnan(ai_pct) else "-",
                        str(hu_pos), f"{hu_pct:.1f}%" if not np.isnan(hu_pct) else "-",
                        f"{reduction:+.1f}%" if not np.isnan(reduction) else "-",
                        lf_cap, eff_score,
                    ])

        if eff_rows:
            tbl = doc.add_table(rows=1 + len(eff_rows), cols=len(eff_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(tbl)
            header_row(tbl, eff_headers)
            for i, row_data in enumerate(eff_rows):
                for j, val in enumerate(row_data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
                    if j == 9 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val >= 95:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val < 80:
                                shade(tbl.cell(i + 1, j), "FFD6D6")
                        except ValueError:
                            pass
                    if j == 5 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val < 30:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val > 70:
                                shade(tbl.cell(i + 1, j), "FFF3CD")
                        except ValueError:
                            pass

            doc.add_paragraph()
            add_note(doc, "AI Positives = articles the AI selected for further review. "
                     "Reduction = % fewer articles selected by AI vs human. Negative = AI selected more. "
                     "Efficiency Score = LF Capture Rate × (1 − AI Positive Rate). Higher = better.")
        else:
            p = doc.add_paragraph("Insufficient data for absolute efficiency analysis.")
            p.runs[0].font.italic = True

        doc.add_paragraph()

    # ==================================================================
    #  APPENDIX A — TIAB FALSE POSITIVES
    # ==================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix A. TIAB False Positives by Run", level=1)
    add_note(doc, "Articles included by the AI but excluded by the human screener at the TIAB stage.")

    fp_results = all_results.get("false_positives", {})
    has_fp_data = any(
        any(
            any(fp_results.get(pn, {}).get(mn, {}).get(tn2, {}).get("fp_titles", [])
                for tn2 in fp_results.get(pn, {}).get(mn, {}))
            for mn in fp_results.get(pn, {})
        )
        for pn in fp_results
    )

    if not has_fp_data:
        p = doc.add_paragraph("No false positive data available.")
        p.runs[0].font.italic = True
    else:
        for pn in sorted(fp_results.keys()):
            proj = projects[pn]
            proj_fp = fp_results[pn]

            all_fp_dict = {}
            run_keys_fp = []
            for mn in sorted(proj_fp.keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(proj_fp[mn].keys()):
                    r = proj_fp[mn][test_num]
                    label = f"{model_name} {test_num}º"
                    run_keys_fp.append((mn, test_num, label))
                    for art in r.get("fp_articles", []):
                        t = art.get("title", "")
                        if t and t not in all_fp_dict:
                            all_fp_dict[t] = art
                    for t in r.get("fp_titles", []):
                        if t and t not in all_fp_dict:
                            all_fp_dict[t] = {"title": t, "abstract": ""}

            if not all_fp_dict:
                continue

            add_heading(doc, f"Project: {proj['name']}", level=2)
            add_note(doc, f"{len(all_fp_dict)} unique article(s) classified as false positive.")

            for idx, title in enumerate(sorted(all_fp_dict.keys()), 1):
                art = all_fp_dict[title]
                abstract_text = art.get("abstract", "") or ""
                if pd.isna(abstract_text):
                    abstract_text = ""

                fp_models = []
                for mn, test_num, label in run_keys_fp:
                    r = proj_fp[mn][test_num]
                    if title in r.get("fp_titles", []):
                        fp_models.append(label)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.space_before = Pt(6)
                p_title.paragraph_format.space_after = Pt(2)
                run_num = p_title.add_run(f"{idx}. ")
                run_num.bold = True
                run_num.font.size = Pt(9)
                run_num.font.name = "Times New Roman"
                run_t = p_title.add_run(str(title))
                run_t.bold = True
                run_t.font.size = Pt(9)
                run_t.font.name = "Times New Roman"

                p_models = doc.add_paragraph()
                p_models.paragraph_format.left_indent = Pt(18)
                p_models.paragraph_format.space_before = Pt(0)
                p_models.paragraph_format.space_after = Pt(2)
                run_lbl = p_models.add_run("False positive in: ")
                run_lbl.bold = True
                run_lbl.font.size = Pt(8)
                run_lbl.font.name = "Times New Roman"
                run_lbl.font.color.rgb = RGBColor(200, 0, 0)
                run_mlist = p_models.add_run(", ".join(fp_models) if fp_models else "—")
                run_mlist.font.size = Pt(8)
                run_mlist.font.name = "Times New Roman"

                if abstract_text.strip():
                    p_abs = doc.add_paragraph()
                    p_abs.paragraph_format.left_indent = Pt(18)
                    p_abs.paragraph_format.space_after = Pt(6)
                    run_ab = p_abs.add_run(str(abstract_text))
                    run_ab.italic = True
                    run_ab.font.size = Pt(8)
                    run_ab.font.name = "Times New Roman"

            doc.add_paragraph()

    # ==================================================================
    #  APPENDIX B — FULLTEXT MISSED ARTICLES
    # ==================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix B. Fulltext Missed Articles (Title, Abstract, and Model)", level=1)
    add_note(doc, "Articles included in the human fulltext review but excluded by the AI in at least one run.")

    ft_app_data = all_results.get("fulltext", {})
    has_missed_abs = any(
        any(
            any(ft_app_data.get(pn, {}).get(mn, {}).get(tn2, {}).get("missed_articles", [])
                for tn2 in ft_app_data.get(pn, {}).get(mn, {}))
            for mn in ft_app_data.get(pn, {}))
        for pn in ft_app_data
    )

    if not has_missed_abs:
        p = doc.add_paragraph("No missed fulltext articles found.")
        p.runs[0].font.italic = True
    else:
        for pn in sorted(ft_app_data.keys()):
            proj = projects[pn]
            proj_ft_app = ft_app_data[pn]

            run_keys_miss = []
            for mn in sorted(proj_ft_app.keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(proj_ft_app[mn].keys()):
                    label = f"{model_name} {test_num}º"
                    run_keys_miss.append((mn, test_num, label))

            seen_titles_app = set()
            unique_missed_app = []
            for mn in sorted(proj_ft_app.keys()):
                for test_num in sorted(proj_ft_app[mn].keys()):
                    r = proj_ft_app[mn][test_num]
                    for art in r.get("missed_articles", []):
                        t = art.get("title", "")
                        if t not in seen_titles_app:
                            seen_titles_app.add(t)
                            unique_missed_app.append(art)

            if not unique_missed_app:
                continue

            add_heading(doc, f"Project: {proj['name']}", level=2)
            add_note(doc, f"{len(unique_missed_app)} unique article(s) missed by the AI.")

            # Matrix table
            tn_num = next_table()
            add_heading(doc, f"Table {tn_num}. Fulltext Articles Missed by AI — {proj['name']}", level=3)
            add_note(doc, "Each column shows a model/test run: ✗ = missed, ✓ = captured.")

            miss_headers = ["#", "Article Title"] + [rk[2] for rk in run_keys_miss]
            n_cols = len(miss_headers)

            miss_tbl = doc.add_table(rows=1, cols=n_cols)
            miss_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(miss_tbl)
            header_row(miss_tbl, miss_headers)

            sorted_missed = sorted(unique_missed_app, key=lambda a: a.get("title", ""))
            for idx, art in enumerate(sorted_missed, 1):
                title_app = art.get("title", "—") or "—"
                row = miss_tbl.add_row()
                set_cell(row.cells[0], str(idx), font_size=Pt(7))
                set_cell(row.cells[1], str(title_app)[:150], font_size=Pt(7),
                         align=WD_ALIGN_PARAGRAPH.LEFT)

                for col_idx, (mn, test_num, _label) in enumerate(run_keys_miss, start=2):
                    r = proj_ft_app[mn][test_num]
                    missed_in_run = r.get("missed_titles", [])
                    if title_app in missed_in_run:
                        set_cell(row.cells[col_idx], "✗", font_size=Pt(7),
                                 color=RGBColor(200, 0, 0))
                        shade(row.cells[col_idx], "FFD6D6")
                    else:
                        set_cell(row.cells[col_idx], "✓", font_size=Pt(7),
                                 color=RGBColor(0, 128, 0))
                        shade(row.cells[col_idx], "D5F5E3")

            doc.add_paragraph()

            # Detailed articles
            add_heading(doc, f"Missed Articles Detail — {proj['name']}", level=3)

            for idx, art in enumerate(sorted_missed, 1):
                title_app = art.get("title", "—") or "—"
                abstract_app = art.get("abstract", "—") or "—"
                if pd.isna(abstract_app):
                    abstract_app = "—"

                missed_by = []
                captured_by = []
                for mn, test_num, label in run_keys_miss:
                    r = proj_ft_app[mn][test_num]
                    if title_app in r.get("missed_titles", []):
                        missed_by.append(label)
                    else:
                        captured_by.append(label)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.space_before = Pt(6)
                p_title.paragraph_format.space_after = Pt(2)
                run_num = p_title.add_run(f"{idx}. ")
                run_num.bold = True
                run_num.font.size = Pt(9)
                run_num.font.name = "Times New Roman"
                run_t = p_title.add_run(str(title_app))
                run_t.bold = True
                run_t.font.size = Pt(9)
                run_t.font.name = "Times New Roman"

                p_models = doc.add_paragraph()
                p_models.paragraph_format.left_indent = Pt(18)
                p_models.paragraph_format.space_before = Pt(0)
                p_models.paragraph_format.space_after = Pt(2)
                run_lbl = p_models.add_run("Missed by: ")
                run_lbl.bold = True
                run_lbl.font.size = Pt(8)
                run_lbl.font.name = "Times New Roman"
                run_lbl.font.color.rgb = RGBColor(200, 0, 0)
                run_mlist = p_models.add_run(", ".join(missed_by) if missed_by else "—")
                run_mlist.font.size = Pt(8)
                run_mlist.font.name = "Times New Roman"
                if captured_by:
                    run_cap = p_models.add_run(f"  |  Captured by: ")
                    run_cap.bold = True
                    run_cap.font.size = Pt(8)
                    run_cap.font.name = "Times New Roman"
                    run_cap.font.color.rgb = RGBColor(0, 128, 0)
                    run_clist = p_models.add_run(", ".join(captured_by))
                    run_clist.font.size = Pt(8)
                    run_clist.font.name = "Times New Roman"

                p_abs = doc.add_paragraph()
                p_abs.paragraph_format.left_indent = Pt(18)
                p_abs.paragraph_format.space_after = Pt(6)
                run_ab = p_abs.add_run(str(abstract_app))
                run_ab.italic = True
                run_ab.font.size = Pt(8)
                run_ab.font.name = "Times New Roman"

            doc.add_paragraph()

    # ==================================================================
    #  SAVE
    # ==================================================================
    ts_file = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = output_dir / f"relatorio_unificado_{ts_file}.docx"
    doc.save(str(docx_path))
    return docx_path
