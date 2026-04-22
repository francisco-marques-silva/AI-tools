"""
docx_helpers.py — Word document formatting helpers.
"""

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def shade(cell, color_hex):
    """Apply background colour to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=Pt(9), font_name="Times New Roman", color=None):
    """Write text into a table cell with styling."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = color
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)


def add_borders(table):
    """Add horizontal borders (no vertical lines) — academic style."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_heading(doc, text, level=2):
    """Add a heading with Times New Roman, black colour."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_note(doc, text):
    """Add a small italic note paragraph."""
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    return p


def header_row(table, headers, row_idx=0, color="D9E2F3"):
    """Style the header row of a table."""
    for j, h in enumerate(headers):
        set_cell(table.cell(row_idx, j), h, bold=True, font_size=Pt(8))
        shade(table.cell(row_idx, j), color)
