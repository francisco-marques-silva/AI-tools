"""
RELATORIO UNIFICADO — Analise multiprojeto de triagem por IA.

Gera um unico documento Word (.docx) consolidando TODAS as analises
de TODOS os projetos e modelos encontrados na pasta input/.

Analises realizadas:
  1. Validacao de dados (correspondencia arquivos ↔ metadados)
  2. Resumo de metadados e custos
  3. Analise diagnostica (sensibilidade, especificidade, etc.)
  4. Verificacao de fulltext (capture rate)
  5. Teste-reteste (reprodutibilidade)
  6. Falsos positivos / falsos negativos
  7. Tabela comparativa entre modelos e projetos
  8. Analise de custo-efetividade

Uso:
    python report/relatorio_unificado.py             (auto-detecta tudo em input/)
    python report/relatorio_unificado.py --input_dir input/
"""

import argparse
import datetime
import re
import sys
from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ------------------------------------------------------------------ paths --
SCRIPT_DIR  = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
INPUT_DIR   = PROJECT_DIR / "input"
OUTPUT_DIR  = PROJECT_DIR / "output"


# ===========================================================================
#  UTILIDADES GERAIS
# ===========================================================================

def load_file(path: str) -> pd.DataFrame:
    ext = Path(path).suffix.lower()
    if ext == ".csv":
        try:
            return pd.read_csv(path, encoding="utf-8")
        except UnicodeDecodeError:
            return pd.read_csv(path, encoding="latin-1")
    elif ext in (".xlsx", ".xls"):
        return pd.read_excel(path)
    else:
        raise ValueError(f"Formato nao suportado: {ext}")


def normalise_columns(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    return df


_HTML_TAG_RE = re.compile(r'<[^>]+')
_NONALNUM_RE = re.compile(r'[^a-z0-9]')


def normalise_title(s) -> str:
    """Key used for pairing articles by title.
    Strips HTML tags, then removes all non-alphanumeric characters (lowercase).
    Handles D<sub>1</sub> vs D(1) vs D1, brackets, extra whitespace, etc.
    """
    if pd.isna(s):
        return ""
    s = _HTML_TAG_RE.sub('', str(s))       # strip <sub>, <sup>, etc.
    return _NONALNUM_RE.sub('', s.lower())  # keep only a-z and 0-9


def normalise_decision(s) -> str:
    if pd.isna(s):
        return ""
    d = str(s).strip().lower()
    if d == "included":
        return "include"
    if d == "excluded":
        return "exclude"
    return d


def binarise_decision(d: str) -> str:
    """include/maybe -> maybe  |  exclude -> exclude"""
    d = normalise_decision(d)
    if d in ("include", "maybe"):
        return "maybe"
    return d


def normalise_model_name(name: str) -> str:
    """Padroniza variantes de nome de modelo: gpt-5-2, gpt-5.2, gpt-5_2 -> gpt-5_2"""
    return name.strip().lower().replace(".", "_").replace("-", "_").replace(" ", "_")


# ---------------------------------------------------------------- formats --

def fmt(v, d=4):
    if isinstance(v, int):
        return str(v)
    if isinstance(v, float):
        if np.isnan(v):
            return "N/A"
        if np.isinf(v):
            return "∞"
        return f"{v:.{d}f}"
    return str(v)


def fmt_pct(v):
    if isinstance(v, float) and not np.isnan(v) and not np.isinf(v):
        return f"{v * 100:.1f}%"
    return "-"


# ===========================================================================
#  DETECCAO E PARSING DE ARQUIVOS
# ===========================================================================

# Padrao do nome de arquivo da IA:  YYYYMMDD - modelo - Xo teste - projeto.xlsx
AI_FILE_PATTERN = re.compile(
    r"^(\d{8})\s*-\s*(.+?)\s*-\s*(\d)[ºo°]\s*teste\s*-\s*(.+)\.xlsx$",
    re.IGNORECASE,
)

# Padrao do arquivo humano:  Projeto - TIAB.xlsx, Projeto - Fulltext.xlsx ou Projeto - Listfinal.xlsx
HUMAN_FILE_PATTERN = re.compile(
    r"^(.+?)\s*-\s*(TIAB|Fulltext|Listfinal)\.xlsx$",
    re.IGNORECASE,
)


def parse_ai_filename(filename: str):
    """Retorna dict {code, model, test_num, project} ou None."""
    m = AI_FILE_PATTERN.match(filename)
    if not m:
        return None
    return {
        "code": m.group(1),
        "model": m.group(2).strip(),
        "model_norm": normalise_model_name(m.group(2).strip()),
        "test_num": int(m.group(3)),
        "project": m.group(4).strip(),
        "project_norm": m.group(4).strip().lower(),
    }


def parse_human_filename(filename: str):
    """Retorna dict {project, type} ou None.  type = 'tiab' | 'fulltext'"""
    m = HUMAN_FILE_PATTERN.match(filename)
    if not m:
        return None
    return {
        "project": m.group(1).strip(),
        "project_norm": m.group(1).strip().lower(),
        "type": m.group(2).strip().lower(),
    }


def scan_input_dir(input_dir: Path):
    """Escaneia input/ e retorna estrutura de dados com todos os arquivos."""
    ai_files = []      # lista de dicts
    human_files = []    # lista de dicts
    metadados_path = None

    for f in sorted(input_dir.iterdir()):
        if not f.is_file():
            continue
        name = f.name

        if name.lower() == "metadata.xlsx":
            metadados_path = f
            continue

        parsed_ai = parse_ai_filename(name)
        if parsed_ai:
            parsed_ai["path"] = f
            ai_files.append(parsed_ai)
            continue

        parsed_human = parse_human_filename(name)
        if parsed_human:
            parsed_human["path"] = f
            human_files.append(parsed_human)
            continue

    return ai_files, human_files, metadados_path


def build_project_structure(ai_files, human_files, metadados_path):
    """
    Organiza os dados em estrutura por projeto:
    {
      project_norm: {
        'name': str,
        'human_tiab': Path | None,
        'human_fulltext': Path | None,
        'models': {
          model_norm: {
            'name': str,
            'tests': {1: {path, code}, 2: {path, code}}
          }
        }
      }
    }
    """
    projects = {}

    for ai in ai_files:
        pn = ai["project_norm"]
        mn = ai["model_norm"]
        if pn not in projects:
            projects[pn] = {
                "name": ai["project"],
                "human_tiab": None,
                "human_fulltext": None,
                "human_listfinal": None,
                "models": {},
            }
        if mn not in projects[pn]["models"]:
            projects[pn]["models"][mn] = {
                "name": ai["model"],
                "tests": {},
            }
        projects[pn]["models"][mn]["tests"][ai["test_num"]] = {
            "path": ai["path"],
            "code": ai["code"],
        }

    for hf in human_files:
        pn = hf["project_norm"]
        if pn not in projects:
            projects[pn] = {
                "name": hf["project"],
                "human_tiab": None,
                "human_fulltext": None,
                "human_listfinal": None,
                "models": {},
            }
        if hf["type"] == "tiab":
            projects[pn]["human_tiab"] = hf["path"]
        elif hf["type"] == "fulltext":
            projects[pn]["human_fulltext"] = hf["path"]
        elif hf["type"] == "listfinal":
            projects[pn]["human_listfinal"] = hf["path"]

    # Metadados
    metadados = None
    if metadados_path and metadados_path.is_file():
        metadados = pd.read_excel(metadados_path)
        # Drop empty rows (trailing NaN rows in the spreadsheet)
        if "project" in metadados.columns:
            metadados = metadados.dropna(subset=["project"]).reset_index(drop=True)

    return projects, metadados


# ===========================================================================
#  ANALISES ESTATISTICAS
# ===========================================================================

def confusion_matrix(ai_decisions, human_decisions):
    """Calcula TP, FP, FN, TN (gold standard = humano)."""
    tp = int(((ai_decisions == "maybe")   & (human_decisions == "maybe")).sum())
    fp = int(((ai_decisions == "maybe")   & (human_decisions == "exclude")).sum())
    fn = int(((ai_decisions == "exclude") & (human_decisions == "maybe")).sum())
    tn = int(((ai_decisions == "exclude") & (human_decisions == "exclude")).sum())
    return tp, fp, fn, tn


def calc_metrics(tp, fp, fn, tn):
    n    = tp + fp + fn + tn
    sens = tp / (tp + fn) if (tp + fn) else float("nan")
    spec = tn / (tn + fp) if (tn + fp) else float("nan")
    ppv  = tp / (tp + fp) if (tp + fp) else float("nan")
    npv  = tn / (tn + fn) if (tn + fn) else float("nan")
    acc  = (tp + tn) / n  if n else float("nan")
    prev = (tp + fn) / n  if n else float("nan")
    f1   = 2 * tp / (2 * tp + fp + fn) if (2 * tp + fp + fn) else float("nan")
    youd = sens + spec - 1 if not (np.isnan(sens) or np.isnan(spec)) else float("nan")
    lr_p = sens / (1 - spec) if (1 - spec) > 0 else float("inf")
    lr_n = (1 - sens) / spec if spec > 0 else float("inf")
    return {
        "N": n,
        "Prevalence": prev,
        "Sensitivity": sens,
        "Specificity": spec,
        "PPV (Precision)": ppv,
        "NPV": npv,
        "Accuracy": acc,
        "F1 Score": f1,
        "LR+": lr_p,
        "LR-": lr_n,
        "Youden Index": youd,
    }


def calc_kappa(tp, fp, fn, tn):
    n = tp + fp + fn + tn
    if n == 0:
        return float("nan"), float("nan"), float("nan"), float("nan"), ""
    po = (tp + tn) / n
    pe = ((tp + fp) * (tp + fn) + (fn + tn) * (fp + tn)) / (n * n)
    if pe == 1:
        k = 1.0
    else:
        k = (po - pe) / (1 - pe)
    if (1 - pe) == 0:
        se = float("nan")
    else:
        se = np.sqrt(pe * (1 - pe) / (n * (1 - pe) ** 2))
    ci_lo = k - 1.96 * se
    ci_hi = k + 1.96 * se
    if k < 0:       interp = "Poor (< 0)"
    elif k < 0.20:  interp = "Slight (0.00–0.20)"
    elif k < 0.40:  interp = "Fair (0.21–0.40)"
    elif k < 0.60:  interp = "Moderate (0.41–0.60)"
    elif k < 0.80:  interp = "Substantial (0.61–0.80)"
    else:            interp = "Almost Perfect (0.81–1.00)"
    return k, se, ci_lo, ci_hi, interp


# ---- Pareamento ----

def do_pairing(ai_path: Path, human_path: Path):
    """
    Pareia IA com humano por titulo normalizado.
    Retorna (df_paired, n_unpaired_ai, n_unpaired_human).
    """
    ai_df = normalise_columns(load_file(str(ai_path)))
    hu_df = normalise_columns(load_file(str(human_path)))

    # Normalizar titulos
    ai_df["_title_key"] = ai_df["title"].apply(normalise_title)
    hu_df["_title_key"] = hu_df["title"].apply(normalise_title)

    # Cumcount para duplicatas
    ai_df["_occ"] = ai_df.groupby("_title_key").cumcount().astype(str)
    hu_df["_occ"] = hu_df.groupby("_title_key").cumcount().astype(str)
    ai_df["_merge_key"] = ai_df["_title_key"] + "__" + ai_df["_occ"]
    hu_df["_merge_key"] = hu_df["_title_key"] + "__" + hu_df["_occ"]

    # Renomear coluna de decisao do humano
    if "decision" in hu_df.columns:
        hu_df = hu_df.rename(columns={"decision": "decision_human"})

    # Garantir coluna abstract no ai_df
    if "abstract" not in ai_df.columns:
        ai_df["abstract"] = ""

    merged = pd.merge(
        ai_df[["title", "abstract", "_merge_key", "screening_decision"]],
        hu_df[["_merge_key", "decision_human"]],
        on="_merge_key",
        how="outer",
    )

    # Contabilizar nao pareados
    n_ai_only = merged["screening_decision"].notna() & merged["decision_human"].isna()
    n_hu_only = merged["decision_human"].notna() & merged["screening_decision"].isna()

    # Filtrar apenas pareados
    paired = merged.dropna(subset=["screening_decision", "decision_human"]).copy()

    # Binarizar
    paired["ai_bin"] = paired["screening_decision"].apply(binarise_decision)
    paired["hu_bin"] = paired["decision_human"].apply(binarise_decision)

    return paired, int(n_ai_only.sum()), int(n_hu_only.sum())


# ---- Diagnostica ----

def run_diagnostic(ai_path: Path, human_tiab_path: Path):
    """
    Roda analise diagnostica completa.
    Retorna dict com resultados.
    """
    paired, n_unpaired_ai, n_unpaired_hu = do_pairing(ai_path, human_tiab_path)

    # Filtrar validos
    valid = paired["ai_bin"].isin(["maybe", "exclude"]) & \
            paired["hu_bin"].isin(["maybe", "exclude"])
    paired = paired[valid].copy()

    if paired.empty:
        return None

    tp, fp, fn, tn = confusion_matrix(paired["ai_bin"], paired["hu_bin"])
    metrics = calc_metrics(tp, fp, fn, tn)
    k, se, ci_lo, ci_hi, interp = calc_kappa(tp, fp, fn, tn)

    # Collect titles of FP and FN articles
    fp_mask = (paired["ai_bin"] == "maybe") & (paired["hu_bin"] == "exclude")
    fn_mask = (paired["ai_bin"] == "exclude") & (paired["hu_bin"] == "maybe")

    fp_titles = paired[fp_mask]["title"].dropna().tolist()
    fn_titles = paired[fn_mask]["title"].dropna().tolist()

    # Collect FP articles with abstract (for appendix)
    fp_articles = (
        paired[fp_mask][["title", "abstract"]].to_dict("records")
    )

    return {
        "n_paired": len(paired),
        "n_unpaired_ai": n_unpaired_ai,
        "n_unpaired_hu": n_unpaired_hu,
        "tp": tp, "fp": fp, "fn": fn, "tn": tn,
        "fp_titles": fp_titles,
        "fp_articles": fp_articles,
        "fn_titles": fn_titles,
        "metrics": metrics,
        "kappa": k, "kappa_se": se,
        "kappa_ci_lo": ci_lo, "kappa_ci_hi": ci_hi,
        "kappa_interp": interp,
    }


# ---- Fulltext Check ----

def run_fulltext_check(ai_path: Path, fulltext_path: Path):
    """
    Verifica se os artigos do fulltext seriam mantidos pela IA.
    Retorna dict com resultados.
    """
    ai_df = normalise_columns(load_file(str(ai_path)))
    ft_df = normalise_columns(load_file(str(fulltext_path)))

    ai_df["_title_key"] = ai_df["title"].apply(normalise_title)
    ft_df["_title_key"] = ft_df["title"].apply(normalise_title)

    # Para cada artigo do fulltext, verificar se a IA incluiria
    results = []
    for _, row in ft_df.iterrows():
        tk = normalise_title(row["title"])
        abstract = "" if pd.isna(row.get("abstract", "")) else str(row.get("abstract", "")).strip()
        match = ai_df[ai_df["_title_key"] == tk]
        if match.empty:
            results.append({"title": row["title"], "abstract": abstract,
                            "found": False, "ai_decision": "not_found"})
        else:
            ai_dec = normalise_decision(match.iloc[0]["screening_decision"])
            captured = ai_dec in ("maybe", "include")
            results.append({
                "title": row["title"],
                "abstract": abstract,
                "found": True,
                "ai_decision": ai_dec,
                "captured": captured,
            })

    df_results = pd.DataFrame(results)
    n_total = len(df_results)
    n_found = df_results["found"].sum()
    n_not_found = n_total - n_found

    found_df = df_results[df_results["found"]].copy()
    if not found_df.empty:
        found_df["captured"] = found_df["captured"].astype(bool)
    n_captured = int(found_df["captured"].sum()) if not found_df.empty else 0
    n_missed = int((~found_df["captured"]).sum()) if not found_df.empty else 0
    capture_rate = n_captured / n_found if n_found > 0 else float("nan")
    miss_rate = n_missed / n_found if n_found > 0 else float("nan")

    missed_titles = found_df[~found_df["captured"]]["title"].tolist() if not found_df.empty else []
    missed_articles = (
        found_df[~found_df["captured"]][["title", "abstract"]].to_dict("records")
        if not found_df.empty else []
    )

    return {
        "n_fulltext": n_total,
        "n_found": int(n_found),
        "n_not_found": int(n_not_found),
        "n_captured": n_captured,
        "n_missed": n_missed,
        "capture_rate": capture_rate,
        "miss_rate": miss_rate,
        "missed_titles": missed_titles,
        "missed_articles": missed_articles,
    }


# ---- Listfinal Check ----

def run_listfinal_check(ai_path: Path, listfinal_path: Path):
    """
    Verifica se os artigos da lista final de incluidos seriam mantidos pela IA
    na triagem TIAB.  Semelhante ao fulltext check, mas usa a lista final
    (gold standard definitivo — artigos que sobreviveram ao full-text reading).
    """
    ai_df = normalise_columns(load_file(str(ai_path)))
    lf_df = normalise_columns(load_file(str(listfinal_path)))

    ai_df["_title_key"] = ai_df["title"].apply(normalise_title)
    lf_df["_title_key"] = lf_df["title"].apply(normalise_title)

    results = []
    for _, row in lf_df.iterrows():
        tk = normalise_title(row["title"])
        abstract = "" if pd.isna(row.get("abstract", "")) else str(row.get("abstract", "")).strip()
        match = ai_df[ai_df["_title_key"] == tk]
        if match.empty:
            results.append({"title": row["title"], "abstract": abstract,
                            "found": False, "ai_decision": "not_found"})
        else:
            ai_dec = normalise_decision(match.iloc[0]["screening_decision"])
            captured = ai_dec in ("maybe", "include")
            results.append({
                "title": row["title"],
                "abstract": abstract,
                "found": True,
                "ai_decision": ai_dec,
                "captured": captured,
            })

    df_results = pd.DataFrame(results)
    n_total = len(df_results)
    n_found = int(df_results["found"].sum())
    n_not_found = n_total - n_found

    found_df = df_results[df_results["found"]].copy()
    if not found_df.empty:
        found_df["captured"] = found_df["captured"].astype(bool)
    n_captured = int(found_df["captured"].sum()) if not found_df.empty else 0
    n_missed = int((~found_df["captured"]).sum()) if not found_df.empty else 0
    capture_rate = n_captured / n_found if n_found > 0 else float("nan")
    miss_rate = n_missed / n_found if n_found > 0 else float("nan")

    missed_titles = found_df[~found_df["captured"]]["title"].tolist() if not found_df.empty else []
    missed_articles = (
        found_df[~found_df["captured"]][["title", "abstract"]].to_dict("records")
        if not found_df.empty else []
    )

    return {
        "n_listfinal": n_total,
        "n_found": n_found,
        "n_not_found": n_not_found,
        "n_captured": n_captured,
        "n_missed": n_missed,
        "capture_rate": capture_rate,
        "miss_rate": miss_rate,
        "missed_titles": missed_titles,
        "missed_articles": missed_articles,
    }


# ---- Teste-Reteste ----

def run_test_retest(path_t1: Path, path_t2: Path):
    """
    Compara teste 1 vs teste 2 do mesmo modelo.
    Retorna dict com resultados.
    """
    t1_df = normalise_columns(load_file(str(path_t1)))
    t2_df = normalise_columns(load_file(str(path_t2)))

    # Merge por titulo
    for df in (t1_df, t2_df):
        df["_title_key"] = df["title"].apply(normalise_title)
        df["_occ"] = df.groupby("_title_key").cumcount().astype(str)
        df["_merge_key"] = df["_title_key"] + "__" + df["_occ"]

    merged = pd.merge(
        t1_df[["title", "_merge_key", "screening_decision"]],
        t2_df[["_merge_key", "screening_decision"]],
        on="_merge_key",
        how="inner",
        suffixes=("_t1", "_t2"),
    )

    n_total = len(merged)

    # Binarizar
    merged["t1_bin"] = merged["screening_decision_t1"].apply(binarise_decision)
    merged["t2_bin"] = merged["screening_decision_t2"].apply(binarise_decision)

    # Concordancia exata (decisao original normalizada)
    merged["t1_orig"] = merged["screening_decision_t1"].apply(normalise_decision)
    merged["t2_orig"] = merged["screening_decision_t2"].apply(normalise_decision)
    exact_match = int((merged["t1_orig"] == merged["t2_orig"]).sum())
    exact_pct = exact_match / n_total if n_total > 0 else float("nan")

    # Concordancia binarizada
    binary_match = int((merged["t1_bin"] == merged["t2_bin"]).sum())
    binary_pct = binary_match / n_total if n_total > 0 else float("nan")

    # Kappa (binarizado)
    tp = int(((merged["t1_bin"] == "maybe")   & (merged["t2_bin"] == "maybe")).sum())
    fp = int(((merged["t1_bin"] == "maybe")   & (merged["t2_bin"] == "exclude")).sum())
    fn = int(((merged["t1_bin"] == "exclude") & (merged["t2_bin"] == "maybe")).sum())
    tn = int(((merged["t1_bin"] == "exclude") & (merged["t2_bin"] == "exclude")).sum())
    k, se, ci_lo, ci_hi, interp = calc_kappa(tp, fp, fn, tn)

    # Discordancias
    disc = merged[merged["t1_bin"] != merged["t2_bin"]]
    disc_titles = disc["title"].tolist()

    return {
        "n_total": n_total,
        "exact_match": exact_match,
        "exact_pct": exact_pct,
        "binary_match": binary_match,
        "binary_pct": binary_pct,
        "tp": tp, "fp": fp, "fn": fn, "tn": tn,
        "kappa": k, "kappa_se": se,
        "kappa_ci_lo": ci_lo, "kappa_ci_hi": ci_hi,
        "kappa_interp": interp,
        "n_discordant": len(disc),
        "disc_titles": disc_titles[:50],  # limitar a 50
    }


# ===========================================================================
#  WORD DOCUMENT HELPERS
# ===========================================================================

def shade(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=Pt(9), font_name="Times New Roman", color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = color
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)


def add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_heading(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_note(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    return p


def header_row(table, headers, row_idx=0, color="D9E2F3"):
    for j, h in enumerate(headers):
        set_cell(table.cell(row_idx, j), h, bold=True, font_size=Pt(8))
        shade(table.cell(row_idx, j), color)


# ===========================================================================
#  GERACAO DO RELATORIO WORD
# ===========================================================================

def generate_report(projects, metadados, all_results, output_dir: Path):
    """Gera o documento Word unificado com todas as analises."""

    doc = Document()

    # Estilo padrao
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    table_counter = [0]   # mutable counter

    def next_table():
        table_counter[0] += 1
        return table_counter[0]

    # ==================================================================
    #  CAPA / TITULO
    # ==================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Unified AI Screening Analysis Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated on: {ts}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Lista de projetos
    proj_names = [projects[pn]["name"] for pn in sorted(projects.keys())]
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Projects analyzed: {', '.join(proj_names)}")
    run.font.size = Pt(10)

    doc.add_page_break()

    # ==================================================================
    #  METHODOLOGICAL NOTES & REPORT GUIDE (before Section 1)
    # ==================================================================
    add_heading(doc, "Methodological Notes & Report Guide", level=1)

    # --- Overview of what follows ---
    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(8)
    run = intro_p.add_run(
        "This report consolidates all analyses comparing AI-based screening decisions "
        "against human reviewer decisions across multiple projects and models. "
        "The sections below are organized progressively: from data validation through "
        "diagnostic performance, reproducibility, error analysis, and cost-efficiency."
    )
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    section_guide = [
        ("Section 1 — Data Validation",
         "Inventories all detected input files (AI results, human TIAB, Fulltext, Listfinal, metadata) "
         "and checks correspondence between AI result files and metadata entries. "
         "Alerts are raised for missing or unmatched data."),
        ("Section 2 — Metadata and Costs",
         "Displays the full execution metadata table (model, parameters, execution time, tokens, cost) "
         "and provides cost summaries grouped by project and by model."),
        ("Section 3 — Diagnostic Analysis (AI vs Human TIAB)",
         "For each project × model × test, computes sensitivity, specificity, PPV, NPV, accuracy, F1, "
         "and Cohen's Kappa by comparing AI screening decisions against the human TIAB reference. "
         "Includes 2×2 confusion matrices."),
        ("Section 4 — Fulltext Verification",
         "Measures the capture rate: the proportion of articles selected for full-text reading "
         "that the AI would have retained during TIAB screening. Missed articles are detailed in the Appendix."),
        ("Section 5 — Listfinal Verification (Gold Standard)",
         "The definitive performance measure. Shows the capture rate over the final included articles "
         "(post full-text reading). A high Listfinal capture rate means the AI would not have missed "
         "truly relevant articles."),
        ("Section 6 — Test-Retest (Reproducibility)",
         "Compares two independent runs of the same model on the same dataset, measuring exact agreement, "
         "binarized agreement, and Cohen's Kappa with 95% CI to assess intra-model reproducibility."),
        ("Section 7 — False Negatives",
         "Lists articles included by the human reviewer but excluded by the AI, grouped by model and test. "
         "These are the most critical errors in high-sensitivity screening."),
        ("Section 8 — False Positives",
         "Lists articles excluded by the human reviewer but included by the AI. "
         "Shows the FP rate over human-excluded articles."),
        ("Section 9 — General Comparative Table",
         "A consolidated table presenting all key metrics (sensitivity, specificity, F1, Kappa, "
         "fulltext capture, Listfinal capture, test-retest Kappa, cost) for every model × project × test."),
        ("Section 10 — Cost-Effectiveness",
         "Relates cost (USD) to mean sensitivity per model. Computes cost per sensitivity point "
         "to identify models with the best cost-benefit ratio."),
        ("Section 11 — Workload Reduction",
         "Compares human screening time (time_human) against AI screening time (time_ia) "
         "per execution and per project. Shows time saved, reduction percentage, and speed factor."),
        ("Section 12 — Absolute Efficiency",
         "Combines selectivity and capture into an efficiency score. "
         "Efficiency Score = Listfinal Capture Rate × (1 − AI Positive Rate). "
         "A higher score means the AI retains relevant articles while reducing overall screening volume."),
        ("Appendix A — TIAB False Positives",
         "Detailed per-article list with title, abstract, and which models flagged each article as a false positive."),
        ("Appendix B — Fulltext Missed Articles",
         "Per-article matrix showing which models missed each fulltext-selected article, "
         "with detailed title and abstract information."),
    ]
    for title, desc in section_guide:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run = p.add_run(desc)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    # --- Methodology details ---
    doc.add_paragraph()
    meth_title = doc.add_paragraph()
    run = meth_title.add_run("Key Methodological Definitions")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    method_notes = [
        ("Binarization", "AI decisions (include, maybe, exclude) are binarized for "
         "diagnostic analysis: include and maybe → positive (passes screening), exclude → negative. "
         "This reflects high-sensitivity screening where uncertain articles should not be discarded."),
        ("Gold Standard Hierarchy", "Three reference levels are used: (1) Human TIAB decisions — "
         "the primary comparison for diagnostic metrics; (2) Fulltext selection — articles that "
         "passed initial screening; (3) Listfinal — articles ultimately included after full-text "
         "reading. The Listfinal capture rate (Section 5) is the definitive performance measure."),
        ("Cohen's Kappa", "Interpretation follows Landis & Koch (1977): < 0 Poor; "
         "0.00–0.20 Slight; 0.21–0.40 Fair; 0.41–0.60 Moderate; "
         "0.61–0.80 Substantial; 0.81–1.00 Almost Perfect."),
        ("Workload Reduction", "Human screening time (time_human) is compared with AI screening "
         "time (time_ia). Speed factor = human time / AI time. "
         "Reduction shows the percentage of time saved by using AI."),
        ("Absolute Efficiency", "Efficiency Score = LF Capture Rate × (1 − AI Positive Rate). "
         "A higher score means the AI retains all relevant articles while selecting fewer overall. "
         "AI Positives are articles the AI marked as include/maybe at the TIAB stage."),
        ("Cost-Effectiveness", "Cost per sensitivity point = average cost / (sensitivity × 100). "
         "Lower is better. Allows comparison of models with different pricing."),
    ]
    for title, text in method_notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    doc.add_page_break()

    # ==================================================================
    #  SECAO 1 — VALIDACAO DOS DADOS
    # ==================================================================
    add_heading(doc, "1. Data Validation", level=1)

    tn = next_table()
    add_heading(doc, f"Table {tn}. Detected Files Inventory", level=2)

    # Contar arquivos
    inventory_rows = []
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        n_models = len(proj["models"])
        n_ai_files = sum(len(m["tests"]) for m in proj["models"].values())
        has_tiab = "Yes" if proj["human_tiab"] else "No"
        has_ft = "Yes" if proj["human_fulltext"] else "No"
        has_lf = "Yes" if proj.get("human_listfinal") else "No"
        lf_n = ""
        if proj.get("human_listfinal"):
            try:
                lf_df = load_file(str(proj["human_listfinal"]))
                lf_n = str(len(lf_df))
            except Exception:
                lf_n = "?"
        inventory_rows.append([
            proj["name"], str(n_models), str(n_ai_files), has_tiab, has_ft, has_lf, lf_n
        ])

    inv_headers = ["Project", "Models", "AI Files", "Human TIAB", "Human Fulltext",
                   "Listfinal", "Listfinal N"]
    tbl = doc.add_table(rows=1 + len(inventory_rows), cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, inv_headers)
    for i, row_data in enumerate(inventory_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
            # Highlight red if missing human TIAB
            if j == 3 and val == "No":
                shade(tbl.cell(i + 1, j), "FFD6D6")
            if j == 4 and val == "No":
                shade(tbl.cell(i + 1, j), "FFF3CD")
            if j == 5 and val == "No":
                shade(tbl.cell(i + 1, j), "FFF3CD")

    doc.add_paragraph()

    # Validacao de correspondencia com metadados
    validation_issues = all_results.get("validation_issues", [])
    if validation_issues:
        add_heading(doc, "Validation Notes", level=3)
        for issue in validation_issues:
            p = doc.add_paragraph(f"• {issue}")
            p.runs[0].font.size = Pt(9)
    else:
        p = doc.add_paragraph("✓ All files were validated successfully.")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()

    # ==================================================================
    #  SECAO 2 — METADADOS E CUSTOS
    # ==================================================================
    if metadados is not None and not metadados.empty:
        add_heading(doc, "2. Metadata and Costs", level=1)

        tn = next_table()
        add_heading(doc, f"Table {tn}. Execution Metadata", level=2)
        add_note(doc, "Details of each execution: model, parameters, time, tokens, and costs (USD).")

        meta_cols = ["Project", "Code", "Model", "Parameters", "Version",
                     "Time", "Tokens In", "Tokens Out", "Cost ($)"]
        meta_rows = []
        for _, row in metadados.iterrows():
            meta_rows.append([
                str(row.get("project", "")),
                str(row.get("code", "")),
                str(row.get("model", "")),
                str(row.get("parameter", "")),
                str(row.get("version", "")),
                str(row.get("time_ia", "")),
                str(int(row["tokens input"])) if pd.notna(row.get("tokens input")) else "-",
                str(int(row["tokens_output"])) if pd.notna(row.get("tokens_output")) else "-",
                f"{row['cost_total']:.2f}" if pd.notna(row.get("cost_total")) else "-",
            ])

        tbl = doc.add_table(rows=1 + len(meta_rows), cols=len(meta_cols))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, meta_cols)
        for i, row_data in enumerate(meta_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)

        doc.add_paragraph()

        # Tabela resumo de custos por projeto
        tn = next_table()
        add_heading(doc, f"Table {tn}. Cost Summary by Project", level=2)

        # Filter out empty rows (NaN project)
        meta_valid = metadados.dropna(subset=["project"])
        cost_summary = meta_valid.groupby("project").agg(
            n_execucoes=("cost_total", "count"),
            custo_total=("cost_total", "sum"),
            custo_medio=("cost_total", "mean"),
            tokens_in_total=("tokens input", "sum"),
            tokens_out_total=("tokens_output", "sum"),
        ).reset_index()

        cost_headers = ["Project", "Executions", "Total Cost ($)", "Avg Cost ($)",
                        "Tokens In (total)", "Tokens Out (total)"]
        tbl = doc.add_table(rows=1 + len(cost_summary), cols=6)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cost_headers)
        for i, (_, row) in enumerate(cost_summary.iterrows()):
            vals = [
                str(row["project"]),
                str(int(row["n_execucoes"])),
                f"{row['custo_total']:.2f}",
                f"{row['custo_medio']:.2f}",
                str(int(row["tokens_in_total"])),
                str(int(row["tokens_out_total"])),
            ]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

        # Custo por modelo (cross-project)
        tn = next_table()
        add_heading(doc, f"Table {tn}. Average Cost per Model", level=2)

        cost_model = metadados.dropna(subset=["project"]).copy()
        cost_model["model_norm"] = cost_model["model"].apply(normalise_model_name)
        cost_by_model = cost_model.groupby("model").agg(
            n=("cost_total", "count"),
            custo_medio=("cost_total", "mean"),
            custo_total=("cost_total", "sum"),
        ).reset_index().sort_values("custo_medio")

        cmod_headers = ["Model", "Executions", "Avg Cost ($)", "Total Cost ($)"]
        tbl = doc.add_table(rows=1 + len(cost_by_model), cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cmod_headers)
        for i, (_, row) in enumerate(cost_by_model.iterrows()):
            vals = [str(row["model"]), str(int(row["n"])),
                    f"{row['custo_medio']:.2f}", f"{row['custo_total']:.2f}"]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECAO 3 — ANALISE DIAGNOSTICA POR PROJETO
    # ==================================================================
    section_num = 3
    add_heading(doc, f"{section_num}. TIAB Agreement Analysis (AI vs Human Screener)", level=1)
    add_note(doc, "Comparison between AI and human screener decisions at the TIAB level. "
             "Note: the human TIAB decision is an intermediate reference, not the final gold standard. "
             "The true performance measure is the fulltext capture rate (Section 4).")

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        proj_diag = all_results.get("diagnostic", {}).get(pn, {})

        if not proj_diag:
            p = doc.add_paragraph(f"Project {proj['name']}: no human reference (TIAB), diagnostic analysis not available.")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True
            continue

        add_heading(doc, f"Project: {proj['name']}", level=2)

        # --- Tabela comparativa de metricas entre modelos ---
        tn = next_table()
        add_heading(doc, f"Table {tn}. Model Comparison — {proj['name']}", level=3)
        add_note(doc, "Diagnostic metrics for each model and test, compared to the human gold standard.")

        comp_headers = ["Model", "Test", "N", "TP", "FP", "FN", "TN",
                        "Sens.", "Spec.", "PPV", "NPV", "Acc.", "F1", "Kappa", "95% CI", "Interpretation"]
        comp_rows = []

        for mn in sorted(proj_diag.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_diag[mn].keys()):
                r = proj_diag[mn][test_num]
                if r is None:
                    continue
                m = r["metrics"]
                comp_rows.append([
                    model_name,
                    f"{test_num}º",
                    str(r["n_paired"]),
                    str(r["tp"]), str(r["fp"]), str(r["fn"]), str(r["tn"]),
                    fmt_pct(m["Sensitivity"]),
                    fmt_pct(m["Specificity"]),
                    fmt_pct(m["PPV (Precision)"]),
                    fmt_pct(m["NPV"]),
                    fmt_pct(m["Accuracy"]),
                    fmt(m["F1 Score"], 3),
                    fmt(r["kappa"], 3),
                    f"[{fmt(r['kappa_ci_lo'], 2)}, {fmt(r['kappa_ci_hi'], 2)}]",
                    r["kappa_interp"],
                ])

        tbl = doc.add_table(rows=1 + len(comp_rows), cols=len(comp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, comp_headers)
        for i, row_data in enumerate(comp_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
                # Highlight sensitivity >= 95%
                if j == 7 and val not in ("-", "N/A"):
                    try:
                        v = float(val.replace("%", ""))
                        if v >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif v < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass

        doc.add_paragraph()

        # --- Matrizes de confusao detalhadas por modelo ---
        for mn in sorted(proj_diag.keys()):
            model_name = proj["models"][mn]["name"]
            for test_num in sorted(proj_diag[mn].keys()):
                r = proj_diag[mn][test_num]
                if r is None:
                    continue

                tn_num = next_table()
                add_heading(doc, f"Table {tn_num}. Confusion Matrix — {model_name} (Test {test_num})", level=3)

                cm_tbl = doc.add_table(rows=4, cols=4)
                cm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                add_borders(cm_tbl)

                cm_headers = ["", "Human: Maybe", "Human: Exclude", "Total"]
                for j, h in enumerate(cm_headers):
                    set_cell(cm_tbl.cell(0, j), h, bold=True, font_size=Pt(8))
                    shade(cm_tbl.cell(0, j), "D9E2F3")

                cm_data = [
                    ("AI: Maybe",   str(r["tp"]), str(r["fp"]), str(r["tp"] + r["fp"])),
                    ("AI: Exclude", str(r["fn"]), str(r["tn"]), str(r["fn"] + r["tn"])),
                    ("Total",       str(r["tp"] + r["fn"]), str(r["fp"] + r["tn"]),
                     str(r["tp"] + r["fp"] + r["fn"] + r["tn"])),
                ]
                for i, (label, *vals) in enumerate(cm_data, start=1):
                    set_cell(cm_tbl.cell(i, 0), label, bold=True,
                             font_size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT)
                    for j2, v in enumerate(vals, start=1):
                        set_cell(cm_tbl.cell(i, j2), v, font_size=Pt(8))
                    # Highlight TP e TN
                    if i == 1:
                        shade(cm_tbl.cell(i, 1), "D5F5E3")  # TP
                    if i == 2:
                        shade(cm_tbl.cell(i, 2), "D5F5E3")  # TN

                doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 4 — FULLTEXT VERIFICATION
    # ==================================================================
    section_num = 4
    add_heading(doc, f"{section_num}. Fulltext Verification (Capture Rate)", level=1)
    add_note(doc, "Checks whether articles included in the final review (fulltext) would have been retained by the AI during TIAB screening.")

    ft_results = all_results.get("fulltext", {})
    if not ft_results:
        p = doc.add_paragraph("No project with fulltext data available.")
        p.runs[0].font.italic = True
    else:
        for pn in sorted(ft_results.keys()):
            proj = projects[pn]
            add_heading(doc, f"Project: {proj['name']}", level=2)

            proj_ft = ft_results[pn]

            # Capture rate comparison table
            tn_num = next_table()
            add_heading(doc, f"Table {tn_num}. Fulltext Capture Rate — {proj['name']}", level=3)

            ft_headers = ["Model", "Test", "FT Articles", "Found",
                          "Captured", "Missed", "Capture Rate", "Miss Rate"]
            ft_rows = []
            for mn in sorted(proj_ft.keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(proj_ft[mn].keys()):
                    r = proj_ft[mn][test_num]
                    ft_rows.append([
                        model_name, f"{test_num}º",
                        str(r["n_fulltext"]), str(r["n_found"]),
                        str(r["n_captured"]), str(r["n_missed"]),
                        fmt_pct(r["capture_rate"]), fmt_pct(r["miss_rate"]),
                    ])

            tbl = doc.add_table(rows=1 + len(ft_rows), cols=len(ft_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(tbl)
            header_row(tbl, ft_headers)
            for i, row_data in enumerate(ft_rows):
                for j, val in enumerate(row_data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                    # Highlight capture rate
                    if j == 6 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val >= 95:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val < 80:
                                shade(tbl.cell(i + 1, j), "FFD6D6")
                        except ValueError:
                            pass
                    # Highlight missed > 0
                    if j == 5:
                        try:
                            if int(val) > 0:
                                shade(tbl.cell(i + 1, j), "FFF3CD")
                        except ValueError:
                            pass

            doc.add_paragraph()

            # Note: detailed per-article missed table moved to Appendix
            all_missed = set()
            for mn in sorted(proj_ft.keys()):
                for test_num in sorted(proj_ft[mn].keys()):
                    r = proj_ft[mn][test_num]
                    for t in r.get("missed_titles", []):
                        all_missed.add(t)
            if all_missed:
                p = doc.add_paragraph(
                    f"Note: {len(all_missed)} unique article(s) were missed by at least one AI run. "
                    "See Appendix (Fulltext Missed Articles) for detailed per-article and per-model information."
                )
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.italic = True

            doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 5 — LISTFINAL VERIFICATION (TRUE GOLD STANDARD)
    # ==================================================================
    section_num = 5
    add_heading(doc, f"{section_num}. Listfinal Verification (True Gold Standard Capture Rate)", level=1)
    add_note(doc, "Checks whether articles in the final included list (Listfinal — post full-text reading) "
             "would have been retained by the AI during TIAB screening. "
             "This is the definitive measure of AI screening performance: "
             "the proportion of truly relevant articles that the AI would not have missed.")

    lf_results = all_results.get("listfinal", {})
    if not lf_results:
        p = doc.add_paragraph("No project with Listfinal data available.")
        p.runs[0].font.italic = True
    else:
        # ---- Summary table across all projects ----
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Listfinal Capture Rate — All Projects", level=2)
        add_note(doc, "How many of the final included articles would have been retained by each AI model at TIAB screening.")

        lf_headers = ["Project", "Model", "Test", "Listfinal N", "Found",
                       "Captured", "Missed", "Capture Rate", "Miss Rate"]
        lf_rows = []
        for pn in sorted(lf_results.keys()):
            proj = projects[pn]
            for mn in sorted(lf_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(lf_results[pn][mn].keys()):
                    r = lf_results[pn][mn][test_num]
                    lf_rows.append([
                        proj["name"], model_name, f"{test_num}o",
                        str(r["n_listfinal"]), str(r["n_found"]),
                        str(r["n_captured"]), str(r["n_missed"]),
                        fmt_pct(r["capture_rate"]), fmt_pct(r["miss_rate"]),
                    ])

        tbl = doc.add_table(rows=1 + len(lf_rows), cols=len(lf_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, lf_headers)
        for i, row_data in enumerate(lf_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                # Highlight capture rate
                if j == 7 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 95:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val < 80:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                    except ValueError:
                        pass
                # Highlight missed > 0
                if j == 6:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFF3CD")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 6 — TEST-RETEST
    # ==================================================================
    section_num = 6
    add_heading(doc, f"{section_num}. Test-Retest (Reproducibility)", level=1)
    add_note(doc, "Compares two runs of the same model on the same dataset. "
             "Evaluates the consistency/reproducibility of AI decisions.")

    tr_results = all_results.get("test_retest", {})
    if not tr_results:
        p = doc.add_paragraph("No test-retest pairs found.")
        p.runs[0].font.italic = True
    else:
        # Tabela comparativa geral de teste-reteste
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Test-Retest Summary — All Projects", level=2)

        tr_headers = ["Project", "Model", "N", "Exact Agree.",
                      "Binary Agree.", "Discordant", "Kappa", "95% CI", "Interpretation"]
        tr_rows = []

        for pn in sorted(tr_results.keys()):
            proj = projects[pn]
            for mn in sorted(tr_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                r = tr_results[pn][mn]
                tr_rows.append([
                    proj["name"], model_name,
                    str(r["n_total"]),
                    f"{r['exact_match']} ({fmt_pct(r['exact_pct'])})",
                    f"{r['binary_match']} ({fmt_pct(r['binary_pct'])})",
                    str(r["n_discordant"]),
                    fmt(r["kappa"], 3),
                    f"[{fmt(r['kappa_ci_lo'], 2)}, {fmt(r['kappa_ci_hi'], 2)}]",
                    r["kappa_interp"],
                ])

        tbl = doc.add_table(rows=1 + len(tr_rows), cols=len(tr_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, tr_headers)
        for i, row_data in enumerate(tr_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)

        doc.add_paragraph()

        # Detail by project
        for pn in sorted(tr_results.keys()):
            proj = projects[pn]
            add_heading(doc, f"Project: {proj['name']}", level=2)

            for mn in sorted(tr_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                r = tr_results[pn][mn]

                tn_num = next_table()
                add_heading(doc, f"Table {tn_num}. Binary Test-Retest Matrix — {model_name} ({proj['name']})", level=3)

                cm_tbl = doc.add_table(rows=4, cols=4)
                cm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                add_borders(cm_tbl)

                cm_headers = ["", "2nd Test: Maybe", "2nd Test: Exclude", "Total"]
                for j, h in enumerate(cm_headers):
                    set_cell(cm_tbl.cell(0, j), h, bold=True, font_size=Pt(8))
                    shade(cm_tbl.cell(0, j), "D9E2F3")

                tp_, fp_, fn_, tn_ = r["tp"], r["fp"], r["fn"], r["tn"]
                cm_data = [
                    ("1st Test: Maybe",   str(tp_), str(fp_), str(tp_ + fp_)),
                    ("1st Test: Exclude", str(fn_), str(tn_), str(fn_ + tn_)),
                    ("Total", str(tp_ + fn_), str(fp_ + tn_), str(r["n_total"])),
                ]
                for i, (label, *vals) in enumerate(cm_data, start=1):
                    set_cell(cm_tbl.cell(i, 0), label, bold=True,
                             font_size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT)
                    for j2, v in enumerate(vals, start=1):
                        set_cell(cm_tbl.cell(i, j2), v, font_size=Pt(8))
                    if i == 1:
                        shade(cm_tbl.cell(i, 1), "D5F5E3")
                    if i == 2:
                        shade(cm_tbl.cell(i, 2), "D5F5E3")

                doc.add_paragraph()

    doc.add_page_break()

    # ==================================================================
    #  SECTION 6 — FALSE NEGATIVES (missed articles vs human)
    # ==================================================================
    section_num = 7
    add_heading(doc, f"{section_num}. False Negatives Analysis", level=1)
    add_note(doc, "Articles included by the human (maybe) but excluded by the AI (exclude). "
             "False negatives are the most critical in systematic review screening.")

    fn_results = all_results.get("false_negatives", {})
    if fn_results:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Negatives Count by Model", level=2)

        fn_headers = ["Project", "Model", "Test", "Total Paired", "False Neg.", "% of Total"]
        fn_rows = []
        for pn in sorted(fn_results.keys()):
            proj = projects[pn]
            for mn in sorted(fn_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(fn_results[pn][mn].keys()):
                    r = fn_results[pn][mn][test_num]
                    fn_count = r["fn"]
                    n_total = r["n_paired"]
                    fn_pct = fn_count / n_total * 100 if n_total > 0 else 0
                    fn_rows.append([
                        proj["name"], model_name, f"{test_num}º",
                        str(n_total), str(fn_count), f"{fn_pct:.1f}%",
                    ])

        tbl = doc.add_table(rows=1 + len(fn_rows), cols=len(fn_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fn_headers)
        for i, row_data in enumerate(fn_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                # Highlight FN > 0
                if j == 4:
                    try:
                        if int(val) > 0:
                            shade(tbl.cell(i + 1, j), "FFD6D6")
                        else:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                    except ValueError:
                        pass

        doc.add_paragraph()
    else:
        p = doc.add_paragraph("No false negative data (requires human reference).")
        p.runs[0].font.italic = True

    # ==================================================================
    #  SECTION 7 — FALSE POSITIVES
    # ==================================================================
    section_num = 8
    add_heading(doc, f"{section_num}. False Positives Analysis", level=1)
    add_note(doc, "Articles excluded by the human (exclude) but included by the AI (maybe). "
             "These impact the review workload.")

    fp_results = all_results.get("false_positives", {})
    if fp_results:
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. False Positives Count by Model", level=2)

        fp_headers = ["Project", "Model", "Test", "Total Paired", "False Pos.", "% of Total",
                      "Human Excl. Articles", "FP Rate (of excl.)"]
        fp_rows = []
        for pn in sorted(fp_results.keys()):
            proj = projects[pn]
            for mn in sorted(fp_results[pn].keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(fp_results[pn][mn].keys()):
                    r = fp_results[pn][mn][test_num]
                    fp_count = r["fp"]
                    n_total = r["n_paired"]
                    n_excl_human = r["fp"] + r["tn"]
                    fp_pct = fp_count / n_total * 100 if n_total > 0 else 0
                    fp_rate = fp_count / n_excl_human * 100 if n_excl_human > 0 else 0
                    fp_rows.append([
                        proj["name"], model_name, f"{test_num}º",
                        str(n_total), str(fp_count), f"{fp_pct:.1f}%",
                        str(n_excl_human), f"{fp_rate:.1f}%",
                    ])

        tbl = doc.add_table(rows=1 + len(fp_rows), cols=len(fp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, fp_headers)
        for i, row_data in enumerate(fp_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()
    else:
        p = doc.add_paragraph("No false positive data (requires human reference).")
        p.runs[0].font.italic = True

    doc.add_page_break()

    # ==================================================================
    #  SECTION 8 — GENERAL COMPARATIVE TABLE
    # ==================================================================
    section_num = 9
    add_heading(doc, f"{section_num}. General Comparative Table", level=1)
    add_note(doc, "Consolidated view of all metrics by project × model × test.")

    diag = all_results.get("diagnostic", {})
    ft = all_results.get("fulltext", {})
    lf = all_results.get("listfinal", {})
    tr = all_results.get("test_retest", {})

    # Montar tabela giga
    tn_num = next_table()
    add_heading(doc, f"Table {tn_num}. General Comparison — Diagnostic Performance and Reproducibility", level=2)

    big_headers = ["Project", "Model", "Test", "Sens.", "Spec.", "F1",
                   "Kappa (diag)", "FT Capture", "LF Capture", "Kappa (T-R)", "Cost ($)"]
    big_rows = []

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        for mn in sorted(proj["models"].keys()):
            model_info = proj["models"][mn]
            model_name = model_info["name"]
            for test_num in sorted(model_info["tests"].keys()):
                row_vals = [proj["name"], model_name, f"{test_num}o"]

                # Diagnostica
                d = diag.get(pn, {}).get(mn, {}).get(test_num)
                if d:
                    row_vals.append(fmt_pct(d["metrics"]["Sensitivity"]))
                    row_vals.append(fmt_pct(d["metrics"]["Specificity"]))
                    row_vals.append(fmt(d["metrics"]["F1 Score"], 3))
                    row_vals.append(fmt(d["kappa"], 3))
                else:
                    row_vals.extend(["-", "-", "-", "-"])

                # Fulltext
                f_res = ft.get(pn, {}).get(mn, {}).get(test_num)
                if f_res:
                    row_vals.append(fmt_pct(f_res["capture_rate"]))
                else:
                    row_vals.append("-")

                # Listfinal
                lf_res = lf.get(pn, {}).get(mn, {}).get(test_num)
                if lf_res:
                    row_vals.append(fmt_pct(lf_res["capture_rate"]))
                else:
                    row_vals.append("-")

                # Test-retest (only for 1st test, since it's per pair)
                tr_res = tr.get(pn, {}).get(mn)
                if tr_res and test_num == 1:
                    row_vals.append(fmt(tr_res["kappa"], 3))
                elif test_num == 2 and tr_res:
                    row_vals.append("↑")  # referencia ao par
                else:
                    row_vals.append("-")

                # Custo (do metadados)
                code = model_info["tests"][test_num]["code"]
                cost_val = "-"
                if metadados is not None:
                    meta_match = metadados[metadados["code"].astype(str) == str(code)]
                    if not meta_match.empty:
                        cost_val = f"{meta_match.iloc[0]['cost_total']:.2f}"
                row_vals.append(cost_val)

                big_rows.append(row_vals)

    tbl = doc.add_table(rows=1 + len(big_rows), cols=len(big_headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(tbl)
    header_row(tbl, big_headers)
    for i, row_data in enumerate(big_rows):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)

    doc.add_paragraph()

    # ==================================================================
    #  SECTION 9 — COST-EFFECTIVENESS ANALYSIS
    # ==================================================================
    if metadados is not None and diag:
        section_num = 10
        add_heading(doc, f"{section_num}. Cost-Effectiveness Analysis", level=1)
        add_note(doc, "Relationship between cost (USD) and diagnostic performance. "
                 "Evaluates the best balance between cost and screening quality.")

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Cost vs. Sensitivity per Model (test average)", level=2)

        cost_eff_headers = ["Model", "Project", "Avg Sens.", "Avg Spec.",
                            "Avg F1", "Avg Cost ($)", "Cost per Sens. point"]
        cost_eff_rows = []

        for pn in sorted(diag.keys()):
            proj = projects[pn]
            for mn in sorted(diag[pn].keys()):
                model_name = proj["models"][mn]["name"]
                tests = diag[pn][mn]
                sens_vals = []
                spec_vals = []
                f1_vals = []
                cost_vals = []
                for tn2, r in tests.items():
                    if r is None:
                        continue
                    s = r["metrics"]["Sensitivity"]
                    if not np.isnan(s):
                        sens_vals.append(s)
                    sp = r["metrics"]["Specificity"]
                    if not np.isnan(sp):
                        spec_vals.append(sp)
                    f1v = r["metrics"]["F1 Score"]
                    if not np.isnan(f1v):
                        f1_vals.append(f1v)
                    code = proj["models"][mn]["tests"][tn2]["code"]
                    if metadados is not None:
                        meta_m = metadados[metadados["code"].astype(str) == str(code)]
                        if not meta_m.empty and pd.notna(meta_m.iloc[0]["cost_total"]):
                            cost_vals.append(meta_m.iloc[0]["cost_total"])

                avg_sens = np.mean(sens_vals) if sens_vals else float("nan")
                avg_spec = np.mean(spec_vals) if spec_vals else float("nan")
                avg_f1 = np.mean(f1_vals) if f1_vals else float("nan")
                avg_cost = np.mean(cost_vals) if cost_vals else float("nan")

                # Cost per sensitivity point
                cost_per_sens = avg_cost / (avg_sens * 100) if (
                    not np.isnan(avg_cost) and not np.isnan(avg_sens) and avg_sens > 0
                ) else float("nan")

                cost_eff_rows.append([
                    model_name, proj["name"],
                    fmt_pct(avg_sens), fmt_pct(avg_spec),
                    fmt(avg_f1, 3),
                    fmt(avg_cost, 2) if not np.isnan(avg_cost) else "-",
                    fmt(cost_per_sens, 3) if not np.isnan(cost_per_sens) else "-",
                ])

        tbl = doc.add_table(rows=1 + len(cost_eff_rows), cols=len(cost_eff_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, cost_eff_headers)
        for i, row_data in enumerate(cost_eff_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)

        doc.add_paragraph()

    # ==================================================================
    #  SECTION 11 — WORKLOAD REDUCTION ANALYSIS
    # ==================================================================
    if metadados is not None:
        section_num = 11
        add_heading(doc, f"{section_num}. Workload Reduction Analysis", level=1)
        add_note(doc, "Compares human screening time (time_human) with AI screening time (time_ia). "
                 "Quantifies the time savings achieved by using AI-assisted screening.")

        # Parse timedelta columns
        def parse_td(val):
            """Convert a timedelta or string representation to total hours."""
            if pd.isna(val):
                return float("nan")
            if isinstance(val, pd.Timedelta):
                return val.total_seconds() / 3600.0
            s = str(val).strip()
            if not s:
                return float("nan")
            # Handle "X days HH:MM:SS" format
            try:
                td = pd.to_timedelta(s)
                return td.total_seconds() / 3600.0
            except Exception:
                return float("nan")

        def fmt_hours(h):
            """Format hours as 'Xh Ym'."""
            if np.isnan(h):
                return "-"
            hr = int(h)
            mn = int((h - hr) * 60)
            if hr > 0:
                return f"{hr}h {mn:02d}m"
            return f"{mn}m"

        meta_work = metadados.copy()
        meta_work["_h_human"] = meta_work["time_human"].apply(parse_td)
        meta_work["_h_ia"] = meta_work["time_ia"].apply(parse_td)

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Workload Reduction — Per Execution", level=2)

        wr_headers = ["Project", "Model", "Code", "Human Time", "AI Time",
                      "Time Saved", "Reduction (%)", "Speed Factor"]
        wr_rows = []
        for _, row in meta_work.iterrows():
            h_human = row["_h_human"]
            h_ia = row["_h_ia"]
            saved = h_human - h_ia if not (np.isnan(h_human) or np.isnan(h_ia)) else float("nan")
            reduction = (saved / h_human * 100) if (not np.isnan(saved) and h_human > 0) else float("nan")
            factor = (h_human / h_ia) if (not np.isnan(h_human) and not np.isnan(h_ia) and h_ia > 0) else float("nan")
            wr_rows.append([
                str(row.get("project", "")),
                str(row.get("model", "")),
                str(row.get("code", "")),
                fmt_hours(h_human),
                fmt_hours(h_ia),
                fmt_hours(saved),
                f"{reduction:.1f}%" if not np.isnan(reduction) else "-",
                f"{factor:.0f}x" if not np.isnan(factor) else "-",
            ])

        tbl = doc.add_table(rows=1 + len(wr_rows), cols=len(wr_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, wr_headers)
        for i, row_data in enumerate(wr_rows):
            for j, val in enumerate(row_data):
                align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), val, font_size=Pt(8), align=align)
                # Highlight reduction
                if j == 6 and val not in ("-", "N/A"):
                    try:
                        pct_val = float(val.replace("%", ""))
                        if pct_val >= 90:
                            shade(tbl.cell(i + 1, j), "D5F5E3")
                        elif pct_val >= 50:
                            shade(tbl.cell(i + 1, j), "E8F5E9")
                    except ValueError:
                        pass

        doc.add_paragraph()

        # Summary by project
        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Workload Reduction Summary by Project", level=2)

        wr_proj = meta_work.groupby("project").agg(
            human_hours=("_h_human", "first"),   # same human time per project
            avg_ia_hours=("_h_ia", "mean"),
            min_ia_hours=("_h_ia", "min"),
            max_ia_hours=("_h_ia", "max"),
            n_runs=("_h_ia", "count"),
        ).reset_index()

        wp_headers = ["Project", "Human Time", "Avg AI Time", "Fastest AI",
                      "Avg Reduction (%)", "Avg Speed Factor"]
        tbl = doc.add_table(rows=1 + len(wr_proj), cols=len(wp_headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(tbl)
        header_row(tbl, wp_headers)
        for i, (_, row) in enumerate(wr_proj.iterrows()):
            h_h = row["human_hours"]
            avg_ia = row["avg_ia_hours"]
            fast = row["min_ia_hours"]
            avg_red = ((h_h - avg_ia) / h_h * 100) if (not np.isnan(h_h) and not np.isnan(avg_ia) and h_h > 0) else float("nan")
            avg_fac = (h_h / avg_ia) if (not np.isnan(h_h) and not np.isnan(avg_ia) and avg_ia > 0) else float("nan")
            vals = [
                str(row["project"]),
                fmt_hours(h_h),
                fmt_hours(avg_ia),
                fmt_hours(fast),
                f"{avg_red:.1f}%" if not np.isnan(avg_red) else "-",
                f"{avg_fac:.0f}x" if not np.isnan(avg_fac) else "-",
            ]
            for j, v in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(tbl.cell(i + 1, j), v, font_size=Pt(8), align=align)

        doc.add_paragraph()

    # ==================================================================
    #  SECTION 12 — ABSOLUTE EFFICIENCY ANALYSIS
    # ==================================================================
    diag_abs = all_results.get("diagnostic", {})
    lf_abs = all_results.get("listfinal", {})
    if diag_abs and lf_abs:
        section_num = 12
        add_heading(doc, f"{section_num}. Absolute Efficiency Analysis (TIAB Reduction vs Listfinal Retention)", level=1)
        add_note(doc, "Measures how much the AI reduces the TIAB workload (fewer articles selected for further review) "
                 "while still retaining the truly relevant articles (Listfinal capture). "
                 "An efficient model selects fewer TIAB articles as positive while maintaining near-perfect Listfinal capture.")

        tn_num = next_table()
        add_heading(doc, f"Table {tn_num}. Absolute Efficiency — TIAB Selection Volume vs Listfinal Retention", level=2)

        eff_headers = ["Project", "Model", "Test", "TIAB N",
                       "AI Positives", "AI Pos. (%)", "Human Positives", "Human Pos. (%)",
                       "Reduction", "LF Capture", "Efficiency Score"]
        eff_rows = []

        for pn in sorted(projects.keys()):
            proj = projects[pn]
            proj_diag = diag_abs.get(pn, {})
            proj_lf = lf_abs.get(pn, {})

            for mn in sorted(proj["models"].keys()):
                model_info = proj["models"][mn]
                model_name = model_info["name"]
                for test_num in sorted(model_info["tests"].keys()):
                    d = proj_diag.get(mn, {}).get(test_num)
                    l = proj_lf.get(mn, {}).get(test_num)
                    if not d:
                        continue

                    n_total = d["n_paired"]
                    ai_pos = d["tp"] + d["fp"]   # AI selected as positive (maybe)
                    hu_pos = d["tp"] + d["fn"]   # Human selected as positive
                    ai_pct = ai_pos / n_total * 100 if n_total > 0 else float("nan")
                    hu_pct = hu_pos / n_total * 100 if n_total > 0 else float("nan")
                    # Reduction = how many fewer articles AI selects vs human
                    reduction = ((hu_pos - ai_pos) / hu_pos * 100) if hu_pos > 0 else float("nan")

                    lf_cap = "-"
                    eff_score = "-"
                    if l:
                        cap_rate = l["capture_rate"]
                        lf_cap = fmt_pct(cap_rate)
                        # Efficiency score = Listfinal capture rate * (1 - AI positive rate)
                        # Higher = better (captures everything while selecting less)
                        if not np.isnan(cap_rate) and not np.isnan(ai_pct):
                            score = cap_rate * (1 - ai_pct / 100)
                            eff_score = fmt(score, 3)

                    eff_rows.append([
                        proj["name"], model_name, f"{test_num}o",
                        str(n_total),
                        str(ai_pos), f"{ai_pct:.1f}%" if not np.isnan(ai_pct) else "-",
                        str(hu_pos), f"{hu_pct:.1f}%" if not np.isnan(hu_pct) else "-",
                        f"{reduction:+.1f}%" if not np.isnan(reduction) else "-",
                        lf_cap, eff_score,
                    ])

        if eff_rows:
            tbl = doc.add_table(rows=1 + len(eff_rows), cols=len(eff_headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(tbl)
            header_row(tbl, eff_headers)
            for i, row_data in enumerate(eff_rows):
                for j, val in enumerate(row_data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell(tbl.cell(i + 1, j), val, font_size=Pt(7), align=align)
                    # Highlight LF capture
                    if j == 9 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val >= 95:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val < 80:
                                shade(tbl.cell(i + 1, j), "FFD6D6")
                        except ValueError:
                            pass
                    # Highlight AI positive rate (lower = better)
                    if j == 5 and val not in ("-", "N/A"):
                        try:
                            pct_val = float(val.replace("%", ""))
                            if pct_val < 30:
                                shade(tbl.cell(i + 1, j), "D5F5E3")
                            elif pct_val > 70:
                                shade(tbl.cell(i + 1, j), "FFF3CD")
                        except ValueError:
                            pass

            doc.add_paragraph()

            # Interpretation note
            add_note(doc, "AI Positives = articles the AI selected for further review (include + maybe). "
                     "Reduction = % fewer articles selected by AI vs human. Negative = AI selected more. "
                     "Efficiency Score = LF Capture Rate × (1 − AI Positive Rate). "
                     "Higher score = better (retains all relevant articles while selecting fewer overall).")
        else:
            p = doc.add_paragraph("Insufficient data for absolute efficiency analysis (requires both diagnostic and Listfinal data).")
            p.runs[0].font.italic = True

        doc.add_paragraph()

    # ==================================================================
    #  APPENDIX — TIAB FALSE POSITIVES: ARTICLES INCLUDED BY AI,
    #             EXCLUDED BY HUMAN SCREENER
    # ==================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix. TIAB False Positives by Run", level=1)
    add_note(doc, "Articles included by the AI (maybe/include) but excluded by the human screener at the TIAB stage. "
             "Each entry shows the article title, abstract, and which model/test classified it as false positive.")

    fp_results = all_results.get("false_positives", {})
    has_fp_data = any(
        any(
            any(fp_results.get(pn, {}).get(mn, {}).get(tn2, {}).get("fp_titles", [])
                for tn2 in fp_results.get(pn, {}).get(mn, {}))
            for mn in fp_results.get(pn, {})
        )
        for pn in fp_results
    )

    if not has_fp_data:
        p = doc.add_paragraph("No false positive data available (requires human TIAB reference).")
        p.runs[0].font.italic = True
    else:
        for pn in sorted(fp_results.keys()):
            proj = projects[pn]
            proj_fp = fp_results[pn]

            # Collect all unique FP articles with abstracts; build run_keys
            all_fp_dict = {}  # title -> {"title": ..., "abstract": ...}
            run_keys_fp = []
            for mn in sorted(proj_fp.keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(proj_fp[mn].keys()):
                    r = proj_fp[mn][test_num]
                    label = f"{model_name} {test_num}º"
                    run_keys_fp.append((mn, test_num, label))
                    for art in r.get("fp_articles", []):
                        t = art.get("title", "")
                        if t and t not in all_fp_dict:
                            all_fp_dict[t] = art
                    # Fallback: if fp_articles not available, use fp_titles
                    for t in r.get("fp_titles", []):
                        if t and t not in all_fp_dict:
                            all_fp_dict[t] = {"title": t, "abstract": ""}

            if not all_fp_dict:
                continue

            add_heading(doc, f"Project: {proj['name']}", level=2)
            add_note(doc, f"{len(all_fp_dict)} unique article(s) classified as false positive in at least one run.")

            for idx, title in enumerate(sorted(all_fp_dict.keys()), 1):
                art = all_fp_dict[title]
                abstract_text = art.get("abstract", "") or ""
                if pd.isna(abstract_text):
                    abstract_text = ""

                # Determine which models flagged this as FP
                fp_models = []
                for mn, test_num, label in run_keys_fp:
                    r = proj_fp[mn][test_num]
                    fp_in_run = r.get("fp_titles", [])
                    if title in fp_in_run:
                        fp_models.append(label)

                # Numbered title
                p_title = doc.add_paragraph()
                p_title.paragraph_format.space_before = Pt(6)
                p_title.paragraph_format.space_after = Pt(2)
                run_num = p_title.add_run(f"{idx}. ")
                run_num.bold = True
                run_num.font.size = Pt(9)
                run_num.font.name = "Times New Roman"
                run_t = p_title.add_run(str(title))
                run_t.bold = True
                run_t.font.size = Pt(9)
                run_t.font.name = "Times New Roman"

                # Models that flagged as FP
                p_models = doc.add_paragraph()
                p_models.paragraph_format.left_indent = Pt(18)
                p_models.paragraph_format.space_before = Pt(0)
                p_models.paragraph_format.space_after = Pt(2)
                run_lbl = p_models.add_run("False positive in: ")
                run_lbl.bold = True
                run_lbl.font.size = Pt(8)
                run_lbl.font.name = "Times New Roman"
                run_lbl.font.color.rgb = RGBColor(200, 0, 0)
                run_mlist = p_models.add_run(", ".join(fp_models) if fp_models else "—")
                run_mlist.font.size = Pt(8)
                run_mlist.font.name = "Times New Roman"

                # Abstract
                if abstract_text.strip():
                    p_abs = doc.add_paragraph()
                    p_abs.paragraph_format.left_indent = Pt(18)
                    p_abs.paragraph_format.space_after = Pt(6)
                    run_ab = p_abs.add_run(str(abstract_text))
                    run_ab.italic = True
                    run_ab.font.size = Pt(8)
                    run_ab.font.name = "Times New Roman"

            doc.add_paragraph()

    # ==================================================================
    #  APPENDIX — FULLTEXT MISSED ARTICLES: TITLE, ABSTRACT, AND
    #             WHICH AI MODELS MISSED THEM
    # ==================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix. Fulltext Missed Articles (Title, Abstract, and Model)", level=1)
    add_note(doc, "Articles included in the human fulltext review but excluded by the AI in at least "
             "one run. For each article: title, abstract, and which model/test runs missed it "
             "(✗ = missed, ✓ = captured).")

    ft_app_data = all_results.get("fulltext", {})
    has_missed_abs = any(
        any(
            any(ft_app_data.get(pn, {}).get(mn, {}).get(tn2, {}).get("missed_articles", [])
                for tn2 in ft_app_data.get(pn, {}).get(mn, {}))
            for mn in ft_app_data.get(pn, {}))
        for pn in ft_app_data
    )

    if not has_missed_abs:
        p = doc.add_paragraph("No missed fulltext articles found.")
        p.runs[0].font.italic = True
    else:
        for pn in sorted(ft_app_data.keys()):
            proj = projects[pn]
            proj_ft_app = ft_app_data[pn]

            # Build run_keys for this project
            run_keys_miss = []
            for mn in sorted(proj_ft_app.keys()):
                model_name = proj["models"][mn]["name"]
                for test_num in sorted(proj_ft_app[mn].keys()):
                    label = f"{model_name} {test_num}º"
                    run_keys_miss.append((mn, test_num, label))

            # Collect all unique missed articles across all runs (deduplicate by title)
            seen_titles_app = set()
            unique_missed_app = []
            for mn in sorted(proj_ft_app.keys()):
                for test_num in sorted(proj_ft_app[mn].keys()):
                    r = proj_ft_app[mn][test_num]
                    for art in r.get("missed_articles", []):
                        t = art.get("title", "")
                        if t not in seen_titles_app:
                            seen_titles_app.add(t)
                            unique_missed_app.append(art)

            if not unique_missed_app:
                continue

            add_heading(doc, f"Project: {proj['name']}", level=2)
            add_note(doc, f"{len(unique_missed_app)} unique article(s) missed by the AI in at least one run.")

            # --- Summary matrix table: which models missed which articles ---
            tn_num = next_table()
            add_heading(doc, f"Table {tn_num}. Fulltext Articles Missed by AI — {proj['name']}", level=3)
            add_note(doc, "Each column shows a model/test run: ✗ = missed, ✓ = captured.")

            miss_headers = ["#", "Article Title"] + [rk[2] for rk in run_keys_miss]
            n_cols = len(miss_headers)

            miss_tbl = doc.add_table(rows=1, cols=n_cols)
            miss_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_borders(miss_tbl)
            header_row(miss_tbl, miss_headers)

            sorted_missed = sorted(unique_missed_app, key=lambda a: a.get("title", ""))
            for idx, art in enumerate(sorted_missed, 1):
                title_app = art.get("title", "—") or "—"
                row = miss_tbl.add_row()
                set_cell(row.cells[0], str(idx), font_size=Pt(7))
                set_cell(row.cells[1], str(title_app)[:150], font_size=Pt(7),
                         align=WD_ALIGN_PARAGRAPH.LEFT)

                for col_idx, (mn, test_num, _label) in enumerate(run_keys_miss, start=2):
                    r = proj_ft_app[mn][test_num]
                    missed_in_run = r.get("missed_titles", [])
                    if title_app in missed_in_run:
                        set_cell(row.cells[col_idx], "✗", font_size=Pt(7),
                                 color=RGBColor(200, 0, 0))
                        shade(row.cells[col_idx], "FFD6D6")
                    else:
                        set_cell(row.cells[col_idx], "✓", font_size=Pt(7),
                                 color=RGBColor(0, 128, 0))
                        shade(row.cells[col_idx], "D5F5E3")

            doc.add_paragraph()

            # --- Detailed articles with abstracts ---
            add_heading(doc, f"Missed Articles Detail — {proj['name']}", level=3)

            for idx, art in enumerate(sorted_missed, 1):
                title_app = art.get("title", "—") or "—"
                abstract_app = art.get("abstract", "—") or "—"
                if pd.isna(abstract_app):
                    abstract_app = "—"

                # Determine which models missed this article
                missed_by = []
                captured_by = []
                for mn, test_num, label in run_keys_miss:
                    r = proj_ft_app[mn][test_num]
                    missed_in_run = r.get("missed_titles", [])
                    if title_app in missed_in_run:
                        missed_by.append(label)
                    else:
                        captured_by.append(label)

                # Numbered title
                p_title = doc.add_paragraph()
                p_title.paragraph_format.space_before = Pt(6)
                p_title.paragraph_format.space_after = Pt(2)
                run_num = p_title.add_run(f"{idx}. ")
                run_num.bold = True
                run_num.font.size = Pt(9)
                run_num.font.name = "Times New Roman"
                run_t = p_title.add_run(str(title_app))
                run_t.bold = True
                run_t.font.size = Pt(9)
                run_t.font.name = "Times New Roman"

                # Models that missed / captured
                p_models = doc.add_paragraph()
                p_models.paragraph_format.left_indent = Pt(18)
                p_models.paragraph_format.space_before = Pt(0)
                p_models.paragraph_format.space_after = Pt(2)
                run_lbl = p_models.add_run("Missed by: ")
                run_lbl.bold = True
                run_lbl.font.size = Pt(8)
                run_lbl.font.name = "Times New Roman"
                run_lbl.font.color.rgb = RGBColor(200, 0, 0)
                run_mlist = p_models.add_run(", ".join(missed_by) if missed_by else "—")
                run_mlist.font.size = Pt(8)
                run_mlist.font.name = "Times New Roman"
                if captured_by:
                    run_cap = p_models.add_run(f"  |  Captured by: ")
                    run_cap.bold = True
                    run_cap.font.size = Pt(8)
                    run_cap.font.name = "Times New Roman"
                    run_cap.font.color.rgb = RGBColor(0, 128, 0)
                    run_clist = p_models.add_run(", ".join(captured_by))
                    run_clist.font.size = Pt(8)
                    run_clist.font.name = "Times New Roman"

                # Abstract
                p_abs = doc.add_paragraph()
                p_abs.paragraph_format.left_indent = Pt(18)
                p_abs.paragraph_format.space_after = Pt(6)
                run_ab = p_abs.add_run(str(abstract_app))
                run_ab.italic = True
                run_ab.font.size = Pt(8)
                run_ab.font.name = "Times New Roman"

            doc.add_paragraph()

    # ==================================================================
    #  SALVAR
    # ==================================================================
    ts_file = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = output_dir / f"relatorio_unificado_{ts_file}.docx"
    doc.save(str(docx_path))
    return docx_path


# ===========================================================================
#  LOGICA PRINCIPAL
# ===========================================================================

def validate_data(projects, metadados):
    """Verifica correspondencia entre arquivos e metadados."""
    issues = []

    for pn, proj in projects.items():
        # Project without human reference
        if not proj["human_tiab"]:
            issues.append(
                f"Project '{proj['name']}': no human TIAB file. "
                "Diagnostic analysis will not be possible (test-retest only)."
            )
        if not proj["human_fulltext"]:
            issues.append(
                f"Project '{proj['name']}': no human Fulltext file. "
                "Fulltext verification will not be possible."
            )
        if not proj.get("human_listfinal"):
            issues.append(
                f"Project '{proj['name']}': no Listfinal file. "
                "Listfinal verification will not be possible."
            )

        # Check if all models have test-retest pairs
        for mn, model in proj["models"].items():
            tests = model["tests"]
            if 1 not in tests or 2 not in tests:
                missing = [t for t in [1, 2] if t not in tests]
                issues.append(
                    f"Project '{proj['name']}', model '{model['name']}': "
                    f"missing test {'st, '.join(str(m) for m in missing)}."
                )

    # Check correspondence with metadata
    if metadados is not None:
        meta_codes = set(metadados["code"].astype(str).tolist())
        ai_codes = set()
        for proj in projects.values():
            for model in proj["models"].values():
                for test in model["tests"].values():
                    ai_codes.add(test["code"])

        missing_in_meta = ai_codes - meta_codes
        missing_in_files = meta_codes - ai_codes

        if missing_in_meta:
            issues.append(
                f"Codes present in AI files but missing from metadata: "
                f"{', '.join(sorted(missing_in_meta))}"
            )
        if missing_in_files:
            issues.append(
                f"Codes present in metadata but without corresponding AI file: "
                f"{', '.join(sorted(missing_in_files))}"
            )

    # Check model name inconsistency between metadata and files
    if metadados is not None:
        for pn, proj in projects.items():
            for mn, model in proj["models"].items():
                for tn2, test in model["tests"].items():
                    code = test["code"]
                    meta_match = metadados[metadados["code"].astype(str) == str(code)]
                    if not meta_match.empty:
                        meta_model = meta_match.iloc[0]["model"]
                        file_model = model["name"]
                        meta_model_norm = normalise_model_name(str(meta_model))
                        file_model_norm = normalise_model_name(file_model)
                        if meta_model_norm != file_model_norm:
                            issues.append(
                                f"Model name inconsistency for code {code}: "
                                f"metadata='{meta_model}', file='{file_model}'"
                            )

    return issues


def run_all_analyses(projects, metadados):
    """Executa todas as analises e retorna dict com resultados."""
    all_results = {}

    # ---- Validacao ----
    issues = validate_data(projects, metadados)
    all_results["validation_issues"] = issues

    print("\n" + "=" * 70)
    print("  UNIFIED REPORT — Processing")
    print("=" * 70)

    if issues:
        print("\n  ⚠ Validation notes:")
        for issue in issues:
            print(f"    • {issue}")
    else:
        print("\n  ✓ All data validated successfully.")

    # ---- Diagnostic analysis ----
    print("\n  Running diagnostic analyses...")
    diag_results = {}
    fn_results = {}
    fp_results = {}

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        if not proj["human_tiab"]:
            print(f"    {proj['name']}: no human TIAB, skipping diagnostic.")
            continue

        diag_results[pn] = {}
        fn_results[pn] = {}
        fp_results[pn] = {}

        for mn in sorted(proj["models"].keys()):
            model = proj["models"][mn]
            diag_results[pn][mn] = {}
            fn_results[pn][mn] = {}
            fp_results[pn][mn] = {}

            for test_num, test_info in sorted(model["tests"].items()):
                print(f"    {proj['name']} / {model['name']} / test {test_num}...", end=" ")
                try:
                    r = run_diagnostic(test_info["path"], proj["human_tiab"])
                    diag_results[pn][mn][test_num] = r
                    if r:
                        fn_results[pn][mn][test_num] = {
                            "fn": r["fn"], "n_paired": r["n_paired"],
                            "fn_titles": r.get("fn_titles", []),
                        }
                        fp_results[pn][mn][test_num] = {
                            "fp": r["fp"], "tn": r["tn"], "n_paired": r["n_paired"],
                            "fp_titles": r.get("fp_titles", []),
                            "fp_articles": r.get("fp_articles", []),
                        }
                        print(f"OK (Sens={fmt_pct(r['metrics']['Sensitivity'])}, "
                              f"Kappa={fmt(r['kappa'], 3)})")
                    else:
                        print("No valid data.")
                except Exception as e:
                    print(f"ERROR: {e}")
                    diag_results[pn][mn][test_num] = None

    all_results["diagnostic"] = diag_results
    all_results["false_negatives"] = fn_results
    all_results["false_positives"] = fp_results

    # ---- Fulltext Check ----
    print("\n  Running fulltext verification...")
    ft_results = {}

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        if not proj["human_fulltext"]:
            print(f"    {proj['name']}: no human fulltext, skipping.")
            continue

        ft_results[pn] = {}
        for mn in sorted(proj["models"].keys()):
            model = proj["models"][mn]
            ft_results[pn][mn] = {}
            for test_num, test_info in sorted(model["tests"].items()):
                print(f"    {proj['name']} / {model['name']} / test {test_num}...", end=" ")
                try:
                    r = run_fulltext_check(test_info["path"], proj["human_fulltext"])
                    ft_results[pn][mn][test_num] = r
                    print(f"OK (Capture={fmt_pct(r['capture_rate'])}, "
                          f"Missed={r['n_missed']})")
                except Exception as e:
                    print(f"ERROR: {e}")

    all_results["fulltext"] = ft_results

    # ---- Listfinal Check ----
    print("\n  Running Listfinal verification...")
    lf_results = {}

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        if not proj.get("human_listfinal"):
            print(f"    {proj['name']}: no Listfinal file, skipping.")
            continue

        lf_results[pn] = {}
        for mn in sorted(proj["models"].keys()):
            model = proj["models"][mn]
            lf_results[pn][mn] = {}
            for test_num, test_info in sorted(model["tests"].items()):
                print(f"    {proj['name']} / {model['name']} / test {test_num}...", end=" ")
                try:
                    r = run_listfinal_check(test_info["path"], proj["human_listfinal"])
                    lf_results[pn][mn][test_num] = r
                    print(f"OK (Capture={fmt_pct(r['capture_rate'])}, "
                          f"Missed={r['n_missed']})")
                except Exception as e:
                    print(f"ERROR: {e}")

    all_results["listfinal"] = lf_results

    # ---- Test-Retest ----
    print("\n  Running test-retest...")
    tr_results = {}

    for pn in sorted(projects.keys()):
        proj = projects[pn]
        tr_results[pn] = {}

        for mn in sorted(proj["models"].keys()):
            model = proj["models"][mn]
            if 1 in model["tests"] and 2 in model["tests"]:
                print(f"    {proj['name']} / {model['name']}...", end=" ")
                try:
                    r = run_test_retest(
                        model["tests"][1]["path"],
                        model["tests"][2]["path"],
                    )
                    tr_results[pn][mn] = r
                    print(f"OK (Agree={fmt_pct(r['binary_pct'])}, "
                          f"Kappa={fmt(r['kappa'], 3)})")
                except Exception as e:
                    print(f"ERROR: {e}")

    all_results["test_retest"] = tr_results

    return all_results


# ===========================================================================
#  MAIN
# ===========================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Generates a unified report of all AI screening analyses.",
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument("--input_dir", "-i", default=None,
                        help="Folder with input files (default: input/).")
    args = parser.parse_args()

    input_dir = Path(args.input_dir) if args.input_dir else INPUT_DIR
    if not input_dir.is_dir():
        print(f"\n  ERROR: Input folder not found: {input_dir}")
        sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ---- Detect files ----
    print(f"\n  Scanning folder: {input_dir}")
    ai_files, human_files, metadados_path = scan_input_dir(input_dir)

    print(f"  AI files found:      {len(ai_files)}")
    print(f"  Human files found:   {len(human_files)}")
    print(f"  Metadata found:      {'Yes' if metadados_path else 'No'}")

    if not ai_files:
        print("\n  ERROR: No AI files found in input/.")
        print("  Expected: YYYYMMDD - model - Xº teste - project.xlsx")
        sys.exit(1)

    # ---- Build structure ----
    projects, metadados = build_project_structure(ai_files, human_files, metadados_path)

    print(f"\n  Projects identified: {len(projects)}")
    for pn in sorted(projects.keys()):
        proj = projects[pn]
        n_models = len(proj["models"])
        models_list = ", ".join(m["name"] for m in proj["models"].values())
        print(f"    • {proj['name']}: {n_models} models ({models_list})")

    # ---- Run analyses ----
    all_results = run_all_analyses(projects, metadados)

    # ---- Generate report ----
    print("\n  Generating Word report...")
    docx_path = generate_report(projects, metadados, all_results, OUTPUT_DIR)

    print(f"\n  ✓ Report generated: {docx_path.name}")
    print(f"    Folder: {OUTPUT_DIR}")
    print("=" * 70)
    print()


if __name__ == "__main__":
    main()
